import io
import re
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

from config.settings import (
    MACFOR_AZUL, MACFOR_AZUL_CLARO, MACFOR_CINZA, MACFOR_VERDE,
    MACFOR_BRANCO, COR_FUNDO_HEADER_TAB, COR_FUNDO_LINHA_ALT
)


_MESES_PT = {
    "January": "Janeiro", "February": "Fevereiro", "March": "Março",
    "April": "Abril", "May": "Maio", "June": "Junho",
    "July": "Julho", "August": "Agosto", "September": "Setembro",
    "October": "Outubro", "November": "Novembro", "December": "Dezembro"
}


def _traduzir_mes(texto):
    for en, pt in _MESES_PT.items():
        texto = texto.replace(en, pt)
    return texto


def _configurar_estilos(doc):
    """Configura os estilos globais do documento."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color, bold) in enumerate([
        (Pt(22), MACFOR_AZUL, True),
        (Pt(16), MACFOR_AZUL, True),
        (Pt(13), MACFOR_AZUL_CLARO, True),
    ], start=1):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = size
        h.font.color.rgb = color
        h.font.bold = bold
        h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        h.paragraph_format.space_after = Pt(8)


def _adicionar_header_footer(doc, mes_ref):
    """Adiciona header e footer profissionais a todas as seções."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1)
        section.footer_distance = Cm(1)
        header = section.header
        header.is_linked_to_previous = False
        htable = header.add_table(rows=1, cols=2, width=Inches(6.5))
        htable.alignment = WD_TABLE_ALIGNMENT.CENTER

        for cell in htable.row_cells(0):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                  '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="1B3A5C"/>'
                                  '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '</w:tcBorders>')
            tcPr.append(tcBorders)

        left_cell = htable.cell(0, 0)
        p = left_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("MACFOR")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = MACFOR_AZUL
        run = p.add_run("  |  Relatório Executivo")
        run.font.size = Pt(9)
        run.font.color.rgb = MACFOR_CINZA

        right_cell = htable.cell(0, 1)
        p = right_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Syngenta  |  {mes_ref}")
        run.font.size = Pt(9)
        run.font.color.rgb = MACFOR_CINZA

        if header.paragraphs and header.paragraphs[0].text == '':
            header.paragraphs[0]._element.getparent().remove(header.paragraphs[0]._element)

        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                         '<w:top w:val="single" w:sz="4" w:space="4" w:color="1B3A5C"/>'
                         '</w:pBdr>')
        pPr.append(pBdr)

        run = p.add_run("Confidencial  |  Macfor Inteligência Digital  |  Página ")
        run.font.size = Pt(8)
        run.font.color.rgb = MACFOR_CINZA

        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')

        run_pg = p.add_run()
        run_pg.font.size = Pt(8)
        run_pg.font.color.rgb = MACFOR_CINZA
        run_pg._r.append(fld_char_begin)
        run_pg._r.append(instr_text)
        run_pg._r.append(fld_char_sep)
        run_pg._r.append(fld_char_end)


def _adicionar_capa(doc, mes_ref):
    """Cria uma capa elegante e minimalista."""
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("RELATÓRIO EXECUTIVO")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = MACFOR_AZUL
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Inteligência de Mercado & Performance Digital")
    run.font.size = Pt(14)
    run.font.color.rgb = MACFOR_CINZA
    run.font.name = 'Calibri'

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("_" * 60)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()

    dados_capa = [
        ("CLIENTE", "Syngenta"),
        ("AGÊNCIA", "Macfor Inteligência Digital"),
        ("PERÍODO", mes_ref),
        ("DATA DE EMISSÃO", _traduzir_mes(datetime.now().strftime("%d de %B de %Y"))),
    ]
    for label, valor in dados_capa:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = MACFOR_AZUL
        run = p.add_run(valor)
        run.font.size = Pt(10)
        run.font.color.rgb = MACFOR_CINZA

    doc.add_page_break()


def _adicionar_sumario(doc):
    """Adiciona sumário (Table of Contents) via campo do Word."""
    doc.add_heading("Sumário", level=1)

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)

    run_placeholder = paragraph.add_run("[Atualize o sumário no Word: clique com botão direito > Atualizar campo]")
    run_placeholder.font.color.rgb = MACFOR_CINZA
    run_placeholder.font.size = Pt(9)
    run_placeholder.font.italic = True

    run2 = paragraph.add_run()
    run2._r.append(fld_char_end)

    doc.add_page_break()


def _adicionar_tabela_metricas(doc, titulo, dados_linhas):
    """Adiciona uma tabela de métricas elegante.

    dados_linhas: list de tuples (metrica, atual, variacao_mes, variacao_ano)
    """
    doc.add_heading(titulo, level=2)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    headers = ["Métrica", "Valor Atual", "Var. MoM", "Var. YoY"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = MACFOR_BRANCO
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COR_FUNDO_HEADER_TAB}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    for idx, (metrica, valor, var_mes, var_ano) in enumerate(dados_linhas):
        row = table.add_row()
        valores = [metrica, valor, var_mes, var_ano]
        for i, val in enumerate(valores):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

            if i >= 2 and isinstance(val, str):
                if val.startswith('+'):
                    run.font.color.rgb = MACFOR_VERDE
                elif val.startswith('-'):
                    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

            if idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COR_FUNDO_LINHA_ALT}"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(3)

    doc.add_paragraph()


def _formatar_variacao(valor):
    """Formata variação com sinal."""
    if valor > 0:
        return f"+{valor:.1f}%"
    elif valor < 0:
        return f"{valor:.1f}%"
    return "0.0%"


def _formatar_numero(valor, prefixo="", sufixo=""):
    """Formata número com separador de milhar."""
    if isinstance(valor, float):
        if abs(valor) < 100:
            return f"{prefixo}{valor:,.2f}{sufixo}"
        return f"{prefixo}{valor:,.0f}{sufixo}"
    return f"{prefixo}{valor:,}{sufixo}"


def _aplicar_formatacao_inline(paragraph, texto):
    """Aplica bold e italic inline no texto."""
    paragraph.clear()
    partes = re.split(r'(\*\*.*?\*\*|\*.*?\*)', texto)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith('*') and parte.endswith('*'):
            run = paragraph.add_run(parte[1:-1])
            run.italic = True
        else:
            paragraph.add_run(parte)


def _markdown_para_docx(doc, texto_md, nivel_base=2):
    """Converte texto Markdown simplificado em parágrafos do docx."""
    if not texto_md:
        return

    linhas = texto_md.split('\n')
    for linha in linhas:
        linha_strip = linha.strip()
        if not linha_strip:
            continue

        if linha_strip.startswith('### '):
            doc.add_heading(linha_strip[4:].strip().strip('*'), level=min(nivel_base + 1, 3))
        elif linha_strip.startswith('## '):
            doc.add_heading(linha_strip[3:].strip().strip('*'), level=nivel_base)
        elif linha_strip.startswith('# '):
            doc.add_heading(linha_strip[2:].strip().strip('*'), level=max(nivel_base - 1, 1))
        elif linha_strip.startswith('---'):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif linha_strip.startswith(('- ', '* ', '• ')):
            texto_item = linha_strip.lstrip('-*• ').strip()
            p = doc.add_paragraph(style='List Bullet')
            _aplicar_formatacao_inline(p, texto_item)
        elif re.match(r'^\d+[\.\)] ', linha_strip):
            texto_item = re.sub(r'^\d+[\.\)] ', '', linha_strip).strip()
            p = doc.add_paragraph(style='List Number')
            _aplicar_formatacao_inline(p, texto_item)
        else:
            p = doc.add_paragraph()
            _aplicar_formatacao_inline(p, linha_strip)


def gerar_docx_relatorio(dados, dados_investimentos, dados_custos, dados_seo,
                          etapa_cenario_atual, etapa_destaques, etapa_midias_pagas,
                          etapa_social, etapa_seo, etapa_aprendizados, etapa_proximos_passos):
    """Gera o relatório executivo completo em DOCX (pipeline de 7 etapas)."""
    doc = Document()
    mes_ref = _traduzir_mes(datetime.now().strftime("%B/%Y"))

    _configurar_estilos(doc)
    _adicionar_capa(doc, mes_ref)
    _adicionar_header_footer(doc, mes_ref)
    _adicionar_sumario(doc)

    doc.add_heading("Comparativos de Performance", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Visão consolidada dos principais indicadores de performance digital, "
                     "com variações mês a mês (MoM) e ano a ano (YoY).")
    run.font.italic = True
    run.font.color.rgb = MACFOR_CINZA

    metricas_perf = [
        ("Investimento", _formatar_numero(dados.get('spend_atual', 0), "R$ "),
         _formatar_variacao(dados.get('var_invest_mes', 0)),
         _formatar_variacao(dados.get('var_invest_ano', 0))),
        ("Sessões", _formatar_numero(dados.get('sess_atual', 0)),
         _formatar_variacao(dados.get('var_sess_mes', 0)),
         _formatar_variacao(dados.get('var_sess_ano', 0))),
        ("Alcance (Reach)", _formatar_numero(dados.get('reach_atual', 0)),
         _formatar_variacao(dados.get('var_reach_mes', 0)),
         _formatar_variacao(dados.get('var_reach_ano', 0))),
        ("Video Thruplays", _formatar_numero(dados.get('vtp_atual', 0)),
         _formatar_variacao(dados.get('var_vtp_mes', 0)),
         _formatar_variacao(dados.get('var_vtp_ano', 0))),
        ("Impressões", _formatar_numero(dados.get('imp_atual', 0)),
         _formatar_variacao(dados.get('var_imp_mes', 0)),
         _formatar_variacao(dados.get('var_imp_ano', 0))),
        ("Cliques", _formatar_numero(dados.get('cli_atual', 0)),
         _formatar_variacao(dados.get('var_cli_mes', 0)),
         _formatar_variacao(dados.get('var_cli_ano', 0))),
        ("Engajamentos", _formatar_numero(dados.get('eng_atual', 0)),
         _formatar_variacao(dados.get('var_eng_mes', 0)),
         _formatar_variacao(dados.get('var_eng_ano', 0))),
        ("CTR", _formatar_numero(dados.get('ctr_atual', 0), sufixo="%"),
         _formatar_variacao(dados.get('var_ctr_mes', 0)),
         _formatar_variacao(dados.get('var_ctr_ano', 0))),
    ]
    _adicionar_tabela_metricas(doc, "Indicadores Gerais", metricas_perf)

    inv_linhas = [
        ("Facebook", _formatar_numero(dados_investimentos.get('fb_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_fb_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_fb_ano', 0))),
        ("Instagram", _formatar_numero(dados_investimentos.get('ig_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_ig_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_ig_ano', 0))),
        ("TikTok", _formatar_numero(dados_investimentos.get('tt_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_tt_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_tt_ano', 0))),
        ("Google Ads", _formatar_numero(dados_investimentos.get('google_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_google_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_google_ano', 0))),
        ("Total", _formatar_numero(dados_investimentos.get('total_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_total_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_total_ano', 0))),
    ]
    _adicionar_tabela_metricas(doc, "Investimentos por Canal", inv_linhas)

    custos_linhas = [
        ("CPC", _formatar_numero(dados_custos.get('cpc_atual', 0), "R$ "),
         _formatar_variacao(dados_custos.get('var_cpc_mes', 0)),
         _formatar_variacao(dados_custos.get('var_cpc_ano', 0))),
        ("CPM", _formatar_numero(dados_custos.get('cpm_atual', 0), "R$ "),
         _formatar_variacao(dados_custos.get('var_cpm_mes', 0)),
         _formatar_variacao(dados_custos.get('var_cpm_ano', 0))),
        ("CPE", _formatar_numero(dados_custos.get('cpe_atual', 0), "R$ "),
         _formatar_variacao(dados_custos.get('var_cpe_mes', 0)),
         _formatar_variacao(dados_custos.get('var_cpe_ano', 0))),
        ("CPV", _formatar_numero(dados_custos.get('cpv_atual', 0), "R$ "),
         _formatar_variacao(dados_custos.get('var_cpv_mes', 0)),
         _formatar_variacao(dados_custos.get('var_cpv_ano', 0))),
    ]
    _adicionar_tabela_metricas(doc, "Indicadores de Custo e Eficiência", custos_linhas)

    seo_linhas = [
        ("Visualizações (Total)", _formatar_numero(dados_seo.get('vis_total_atual', 0)),
         _formatar_variacao(dados_seo.get('var_vis_total_mes', 0)),
         _formatar_variacao(dados_seo.get('var_vis_total_ano', 0))),
        ("Sessões Orgânicas", _formatar_numero(dados_seo.get('sess_org_atual', 0)),
         _formatar_variacao(dados_seo.get('var_sess_org_mes', 0)),
         _formatar_variacao(dados_seo.get('var_sess_org_ano', 0))),
        ("Visualizações Orgânicas", _formatar_numero(dados_seo.get('vis_org_atual', 0)),
         _formatar_variacao(dados_seo.get('var_vis_org_mes', 0)),
         _formatar_variacao(dados_seo.get('var_vis_org_ano', 0))),
    ]
    _adicionar_tabela_metricas(doc, "Indicadores Orgânicos", seo_linhas)

    doc.add_page_break()

    doc.add_heading("Cenário Atual", level=1)
    _markdown_para_docx(doc, etapa_cenario_atual)
    doc.add_page_break()

    doc.add_heading("Destaques do Período", level=1)
    _markdown_para_docx(doc, etapa_destaques)
    doc.add_page_break()

    doc.add_heading("Mídias Pagas", level=1)
    _markdown_para_docx(doc, etapa_midias_pagas)
    doc.add_page_break()

    doc.add_heading("Social", level=1)
    _markdown_para_docx(doc, etapa_social)
    doc.add_page_break()

    doc.add_heading("SEO", level=1)
    _markdown_para_docx(doc, etapa_seo)
    doc.add_page_break()

    doc.add_heading("Aprendizados", level=1)
    _markdown_para_docx(doc, etapa_aprendizados)
    doc.add_page_break()

    doc.add_heading("Próximos Passos", level=1)
    _markdown_para_docx(doc, etapa_proximos_passos)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Relatório gerado por IA em {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = MACFOR_CINZA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Macfor Inteligência Digital | macfor.com.br")
    run.font.size = Pt(8)
    run.font.color.rgb = MACFOR_AZUL

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def gerar_docx_cliente(dados, dados_investimentos, dados_custos, dados_seo,
                        etapa_cenario_atual, etapa_destaques, etapa_midias_pagas,
                        etapa_social, etapa_seo, etapa_aprendizados, etapa_proximos_passos):
    """Gera DOCX do relatório para o cliente — narrativo, elegante, focado em valor."""
    doc = Document()
    mes_ref = _traduzir_mes(datetime.now().strftime("%B/%Y"))

    _configurar_estilos(doc)

    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("RELATÓRIO DE RESULTADOS")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = MACFOR_AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Performance Digital & Inteligência Estratégica")
    run.font.size = Pt(14)
    run.font.color.rgb = MACFOR_CINZA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("_" * 60)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()

    for label, valor in [("PREPARADO PARA", "Syngenta"), ("POR", "Macfor Inteligência Digital"), ("PERÍODO", mes_ref)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = MACFOR_AZUL
        run = p.add_run(valor)
        run.font.size = Pt(10)
        run.font.color.rgb = MACFOR_CINZA

    doc.add_page_break()
    _adicionar_header_footer(doc, mes_ref)
    _adicionar_sumario(doc)

    doc.add_heading("Painel de Performance", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Resultados consolidados do período com variações históricas.")
    run.font.italic = True
    run.font.color.rgb = MACFOR_CINZA

    metricas_perf = [
        ("Investimento", _formatar_numero(dados.get('spend_atual', 0), "R$ "),
         _formatar_variacao(dados.get('var_invest_mes', 0)), _formatar_variacao(dados.get('var_invest_ano', 0))),
        ("Alcance", _formatar_numero(dados.get('reach_atual', 0)),
         _formatar_variacao(dados.get('var_reach_mes', 0)), _formatar_variacao(dados.get('var_reach_ano', 0))),
        ("Impressões", _formatar_numero(dados.get('imp_atual', 0)),
         _formatar_variacao(dados.get('var_imp_mes', 0)), _formatar_variacao(dados.get('var_imp_ano', 0))),
        ("Cliques", _formatar_numero(dados.get('cli_atual', 0)),
         _formatar_variacao(dados.get('var_cli_mes', 0)), _formatar_variacao(dados.get('var_cli_ano', 0))),
        ("Engajamentos", _formatar_numero(dados.get('eng_atual', 0)),
         _formatar_variacao(dados.get('var_eng_mes', 0)), _formatar_variacao(dados.get('var_eng_ano', 0))),
        ("CTR", _formatar_numero(dados.get('ctr_atual', 0), sufixo="%"),
         _formatar_variacao(dados.get('var_ctr_mes', 0)), _formatar_variacao(dados.get('var_ctr_ano', 0))),
    ]
    _adicionar_tabela_metricas(doc, "Indicadores de Performance", metricas_perf)
    doc.add_page_break()

    doc.add_heading("Cenário Atual", level=1)
    _markdown_para_docx(doc, etapa_cenario_atual)
    doc.add_page_break()

    doc.add_heading("Destaques do Período", level=1)
    _markdown_para_docx(doc, etapa_destaques)
    doc.add_page_break()

    doc.add_heading("Mídias Pagas", level=1)
    _markdown_para_docx(doc, etapa_midias_pagas)
    doc.add_page_break()

    doc.add_heading("Social", level=1)
    _markdown_para_docx(doc, etapa_social)
    doc.add_page_break()

    doc.add_heading("SEO", level=1)
    _markdown_para_docx(doc, etapa_seo)
    doc.add_page_break()

    doc.add_heading("Aprendizados", level=1)
    _markdown_para_docx(doc, etapa_aprendizados)
    doc.add_page_break()

    doc.add_heading("Próximos Passos", level=1)
    _markdown_para_docx(doc, etapa_proximos_passos)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Macfor Inteligência Digital | macfor.com.br")
    run.font.size = Pt(8)
    run.font.color.rgb = MACFOR_AZUL

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer
