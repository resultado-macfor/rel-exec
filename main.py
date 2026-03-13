import streamlit as st
import pandas as pd
import google.generativeai as genai
import os
import re
from datetime import datetime
from PIL import Image
import base64
import io
from google.cloud import bigquery
from google.oauth2 import service_account
from anthropic import Anthropic
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import plotly.graph_objects as go

# =============================================================================
# CONFIGURAÇÃO INICIAL (DEVE SER A PRIMEIRA COISA DO SCRIPT)
# =============================================================================

st.set_page_config(layout="wide", page_title="Relatório Executivo - IA", page_icon="📊")

# =============================================================================
# LOCK DE SENHA
# =============================================================================
if 'autenticado' not in st.session_state:
    st.session_state.autenticado = False

if not st.session_state.autenticado:
    st.markdown("### 🔒 Acesso Restrito")
    senha_digitada = st.text_input("Digite a senha de acesso:", type="password")
    if st.button("Entrar"):
        if senha_digitada == st.secrets.get("senha_per", ""):
            st.session_state.autenticado = True
            st.rerun()
        else:
            st.error("Senha incorreta.")
    st.stop()

# Inicialização do estado da sessão para evitar KeyError/AttributeError
chaves_sessao = [
    'relatorio_gerado', 'descricoes_imagens', 'descricoes_imagens_mes_passado',
    'dados_processados', 'resumos_social_csvs', 'resumos_seo_csvs',
    'etapa_cenario_atual', 'etapa_destaques', 'etapa_midias_pagas',
    'etapa_social', 'etapa_seo', 'etapa_aprendizados', 'etapa_proximos_passos'
]

for chave in chaves_sessao:
    if chave not in st.session_state:
        if chave == 'relatorio_gerado':
            st.session_state[chave] = False
        elif chave in ['descricoes_imagens', 'descricoes_imagens_mes_passado', 'resumos_social_csvs', 'resumos_seo_csvs']:
            st.session_state[chave] = []
        elif chave == 'dados_processados':
            st.session_state[chave] = {}
        else:
            st.session_state[chave] = ""

# =============================================================================
# CONEXÃO BIGQUERY
# =============================================================================

@st.cache_resource
def get_bigquery_client():
    try:
        service_account_info = None
        
        # 1. Tenta via Streamlit Secrets
        if hasattr(st, 'secrets') and 'gcp_service_account' in st.secrets:
            service_account_info = dict(st.secrets["gcp_service_account"])
            if isinstance(service_account_info.get("private_key"), str):
                service_account_info["private_key"] = service_account_info["private_key"].replace("\\n", "\n")
        
        # 2. Tenta via Variáveis de Ambiente
        elif all(key in os.environ for key in ['project_id', 'private_key', 'client_email']):
            service_account_info = {
                "type": "service_account",
                "project_id": os.environ['project_id'],
                "private_key": os.environ['private_key'].replace('\\n', '\n'),
                "client_email": os.environ['client_email'],
                "token_uri": "https://oauth2.googleapis.com/token",
            }
        
        if not service_account_info:
            return None
        
        credentials = service_account.Credentials.from_service_account_info(service_account_info)
        return bigquery.Client(credentials=credentials, project=service_account_info["project_id"])
    
    except Exception as e:
        st.error(f"❌ Erro na conexão BigQuery: {str(e)}")
        return None

client_bq = get_bigquery_client()
def check_datasources():
    query_check = """
    SELECT DISTINCT datasource 
    FROM `macfor-media-flow.ads.app_view_campaigns`
    WHERE UPPER(account_name) LIKE '%SYNGENTA%'
    """
    try:
        df_sources = client_bq.query(query_check).to_dataframe()
        print("\n" + "="*30)
        print("FONTES ENCONTRADAS (datasource):")
        print(df_sources['datasource'].tolist())
        print("="*30 + "\n")
    except Exception as e:
        print(f"Erro ao verificar fontes: {e}")

# Chame a função para ver no terminal
check_datasources()

def fetch_bigquery_data():
    if client_bq is None:
        st.error("Conexão com BigQuery não disponível.")
        return None
    
    query_expandida = """
    SELECT 
        -- 
        -- INVESTIMENTO POR FONTE (FACEBOOK)
        SUM(CASE WHEN datasource = 'facebook' AND DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN spend ELSE 0 END) as spend_fb_atual,
        SUM(CASE WHEN datasource = 'facebook' AND DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN spend ELSE 0 END) as spend_fb_mes,
        SUM(CASE WHEN datasource = 'facebook' AND DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN spend ELSE 0 END) as spend_fb_ano,

        -- INVESTIMENTO POR FONTE (GOOGLE ADS)
        SUM(CASE WHEN datasource = 'google_ads' AND DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN spend ELSE 0 END) as spend_google_atual,
        SUM(CASE WHEN datasource = 'google_ads' AND DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN spend ELSE 0 END) as spend_google_mes,
        SUM(CASE WHEN datasource = 'google_ads' AND DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN spend ELSE 0 END) as spend_google_ano,

        -- INVESTIMENTO POR FONTE (TIKTOK)
        SUM(CASE WHEN datasource = 'tiktok' AND DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN spend ELSE 0 END) as spend_tiktok_atual,
        SUM(CASE WHEN datasource = 'tiktok' AND DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN spend ELSE 0 END) as spend_tiktok_mes,
        SUM(CASE WHEN datasource = 'tiktok' AND DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN spend ELSE 0 END) as spend_tiktok_ano,


        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN spend ELSE 0 END) as spend_atual,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN spend ELSE 0 END) as spend_mes,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN spend ELSE 0 END) as spend_ano,
        
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN clicks ELSE 0 END) as cli_atual,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN clicks ELSE 0 END) as cli_mes,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN clicks ELSE 0 END) as cli_ano,

        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN actions_page_engagement ELSE 0 END) as eng_atual,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN actions_page_engagement ELSE 0 END) as eng_mes,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN actions_page_engagement ELSE 0 END) as eng_ano,

        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN impressions ELSE 0 END) as imp_atual,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN impressions ELSE 0 END) as imp_mes,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN impressions ELSE 0 END) as imp_ano,

        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN sessions ELSE 0 END) as sess_atual,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN sessions ELSE 0 END) as sess_mes,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN sessions ELSE 0 END) as sess_ano,

        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN reach ELSE 0 END) as reach_atual,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN reach ELSE 0 END) as reach_mes,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN reach ELSE 0 END) as reach_ano,

        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN video_thruplay ELSE 0 END) as vtp_atual,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN video_thruplay ELSE 0 END) as vtp_mes,
        SUM(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN video_thruplay ELSE 0 END) as vtp_ano,

        -- AVGs
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN ctr ELSE NULL END) as ctr_atual,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN ctr ELSE NULL END) as ctr_mes,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN ctr ELSE NULL END) as ctr_ano,

        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN cpc ELSE NULL END) as cpc_atual,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN cpc ELSE NULL END) as cpc_mes,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN cpc ELSE NULL END) as cpc_ano,

        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN cpm ELSE NULL END) as cpm_atual,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN cpm ELSE NULL END) as cpm_mes,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN cpm ELSE NULL END) as cpm_ano,

        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN avg_session_duration ELSE NULL END) as dur_atual,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN avg_session_duration ELSE NULL END) as dur_mes,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN avg_session_duration ELSE NULL END) as dur_ano,

        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN video_thruplay_rate ELSE NULL END) as vtpr_atual,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN video_thruplay_rate ELSE NULL END) as vtpr_mes,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN video_thruplay_rate ELSE NULL END) as vtpr_ano,

        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(CURRENT_DATE(), MONTH) THEN video_thruplay_cpv ELSE NULL END) as vcpv_atual,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 MONTH), MONTH) AND DATE(date) < DATE_TRUNC(CURRENT_DATE(), MONTH) THEN video_thruplay_cpv ELSE NULL END) as vcpv_mes,
        AVG(CASE WHEN DATE(date) >= DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH) AND DATE(date) < DATE_ADD(DATE_TRUNC(DATE_SUB(CURRENT_DATE(), INTERVAL 1 YEAR), MONTH), INTERVAL 1 MONTH) THEN video_thruplay_cpv ELSE NULL END) as vcpv_ano
        
    FROM `macfor-media-flow.ads.app_view_campaigns`
    WHERE UPPER(account_name) LIKE '%SYNGENTA%'
    """
    
    try:
        df = client_bq.query(query_expandida).to_dataframe()
        if not df.empty:
            dados_dict = df.iloc[0].to_dict()
            
            # --- IMPRESSÃO NO TERMINAL (LOG) ---
            print("\n" + "="*60)
            print(f"📊 DADOS RECUPERADOS DO BIGQUERY - {datetime.now().strftime('%H:%M:%S')}")
            print("="*60)
            for chave, valor in dados_dict.items():
                # Formata a exibição: se for número, limita casas decimais
                display_val = f"{valor:,.2f}" if isinstance(valor, (float, int)) else valor
                print(f"{chave.ljust(30)}: {display_val}")
            print("="*60 + "\n")
            
            return dados_dict
        return None
    except Exception as e:
        st.error(f"Erro na query: {str(e)}")
        return None

    




# =============================================================================
# INICIALIZAÇÃO DOS MODELOS
# =============================================================================

# Gemini (sempre usado para visão/imagens + fallback)
gemini_api_key = os.getenv("GEM_API_KEY") or st.secrets.get("GEM_API_KEY", "")
genai.configure(api_key=gemini_api_key)
modelo_gemini = genai.GenerativeModel("gemini-2.5-flash")
modelo_visao = genai.GenerativeModel("gemini-2.5-flash")

# Anthropic (opção para geração de texto)
anthropic_api_key = None
if hasattr(st, 'secrets') and 'ANTH_KEY' in st.secrets:
    anthropic_api_key = st.secrets["ANTH_KEY"]
elif os.getenv("ANTH_KEY"):
    anthropic_api_key = os.getenv("ANTH_KEY")

cliente_anthropic = Anthropic(api_key=anthropic_api_key) if anthropic_api_key else None

def gerar_texto(prompt, modelo_escolhido="Gemini"):
    """Roteia a geração de texto para o modelo escolhido. Usa Gemini como fallback silencioso."""
    if modelo_escolhido == "Claude (Anthropic)" and cliente_anthropic:
        try:
            response = cliente_anthropic.messages.create(
                model="claude-sonnet-4-20250514",
                max_tokens=8096,
                messages=[{"role": "user", "content": prompt}]
            )
            return response.content[0].text
        except Exception:
            # Fallback silencioso para Gemini
            response = modelo_gemini.generate_content(prompt)
            return response.text
    else:
        response = modelo_gemini.generate_content(prompt)
        return response.text

# =============================================================================
# GERAÇÃO DO RELATÓRIO EM DOCX
# =============================================================================

MACFOR_AZUL = RGBColor(0x1B, 0x3A, 0x5C)       # Azul escuro corporativo
MACFOR_AZUL_CLARO = RGBColor(0x2E, 0x86, 0xC1)  # Azul para destaques
MACFOR_CINZA = RGBColor(0x5D, 0x6D, 0x7E)       # Cinza para texto secundário
MACFOR_VERDE = RGBColor(0x27, 0xAE, 0x60)       # Verde para positivo
MACFOR_BRANCO = RGBColor(0xFF, 0xFF, 0xFF)
COR_FUNDO_HEADER_TAB = "1B3A5C"
COR_FUNDO_LINHA_ALT = "EBF5FB"


def _configurar_estilos(doc):
    """Configura os estilos globais do documento."""
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x2C, 0x3E, 0x50)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.15

    for level, (size, color, bold) in enumerate([
        (Pt(22), MACFOR_AZUL, True),
        (Pt(16), MACFOR_AZUL, True),
        (Pt(13), MACFOR_AZUL_CLARO, True),
    ], start=1):
        h = doc.styles[f'Heading {level}']
        h.font.name = 'Calibri'
        h.font.size = size
        h.font.color.rgb = color
        h.font.bold = bold
        h.paragraph_format.space_before = Pt(18 if level == 1 else 14)
        h.paragraph_format.space_after = Pt(8)


def _adicionar_header_footer(doc, mes_ref):
    """Adiciona header e footer profissionais a todas as seções."""
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
        section.header_distance = Cm(1)
        section.footer_distance = Cm(1)

        # --- HEADER ---
        header = section.header
        header.is_linked_to_previous = False
        htable = header.add_table(rows=1, cols=2, width=Inches(6.5))
        htable.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Remove bordas da tabela do header
        for cell in htable.row_cells(0):
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}>'
                                  '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '<w:bottom w:val="single" w:sz="8" w:space="0" w:color="1B3A5C"/>'
                                  '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
                                  '</w:tcBorders>')
            tcPr.append(tcBorders)

        left_cell = htable.cell(0, 0)
        p = left_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run("MACFOR")
        run.font.size = Pt(9)
        run.font.bold = True
        run.font.color.rgb = MACFOR_AZUL
        run = p.add_run("  |  Relatório Executivo")
        run.font.size = Pt(9)
        run.font.color.rgb = MACFOR_CINZA

        right_cell = htable.cell(0, 1)
        p = right_cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run(f"Syngenta  |  {mes_ref}")
        run.font.size = Pt(9)
        run.font.color.rgb = MACFOR_CINZA

        # Remove parágrafo padrão vazio do header
        if header.paragraphs and header.paragraphs[0].text == '':
            header.paragraphs[0]._element.getparent().remove(header.paragraphs[0]._element)

        # --- FOOTER ---
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Linha separadora acima do footer
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>'
                         '<w:top w:val="single" w:sz="4" w:space="4" w:color="1B3A5C"/>'
                         '</w:pBdr>')
        pPr.append(pBdr)

        run = p.add_run("Confidencial  |  Macfor Inteligência Digital  |  Página ")
        run.font.size = Pt(8)
        run.font.color.rgb = MACFOR_CINZA

        # Campo de número de página
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')

        run_pg = p.add_run()
        run_pg.font.size = Pt(8)
        run_pg.font.color.rgb = MACFOR_CINZA
        run_pg._r.append(fld_char_begin)
        run_pg._r.append(instr_text)
        run_pg._r.append(fld_char_sep)
        run_pg._r.append(fld_char_end)


def _adicionar_capa(doc, mes_ref):
    """Cria uma capa elegante e minimalista."""
    # Espaçamento superior
    for _ in range(6):
        doc.add_paragraph()

    # Linha decorativa superior
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(10)

    # Título principal
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("RELATÓRIO EXECUTIVO")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = MACFOR_AZUL
    run.font.name = 'Calibri'

    # Subtítulo
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Inteligência de Mercado & Performance Digital")
    run.font.size = Pt(14)
    run.font.color.rgb = MACFOR_CINZA
    run.font.name = 'Calibri'

    # Linha decorativa inferior
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("_" * 60)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(10)

    # Espaçamento
    for _ in range(3):
        doc.add_paragraph()

    # Bloco Cliente / Agência
    dados_capa = [
        ("CLIENTE", "Syngenta"),
        ("AGÊNCIA", "Macfor Inteligência Digital"),
        ("PERÍODO", mes_ref),
        ("DATA DE EMISSÃO", datetime.now().strftime("%d de %B de %Y").replace(
            "January", "Janeiro").replace("February", "Fevereiro").replace(
            "March", "Março").replace("April", "Abril").replace(
            "May", "Maio").replace("June", "Junho").replace(
            "July", "Julho").replace("August", "Agosto").replace(
            "September", "Setembro").replace("October", "Outubro").replace(
            "November", "Novembro").replace("December", "Dezembro")),
    ]
    for label, valor in dados_capa:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = MACFOR_AZUL
        run = p.add_run(valor)
        run.font.size = Pt(10)
        run.font.color.rgb = MACFOR_CINZA

    # Quebra de página após a capa
    doc.add_page_break()


def _adicionar_sumario(doc):
    """Adiciona sumário (Table of Contents) via campo do Word."""
    p = doc.add_heading("Sumário", level=1)

    # Instrução de campo TOC
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> TOC \\o "1-3" \\h \\z \\u </w:instrText>')
    fld_char_sep = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>')
    fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')

    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)

    # Texto placeholder
    run_placeholder = paragraph.add_run("[Atualize o sumário no Word: clique com botão direito > Atualizar campo]")
    run_placeholder.font.color.rgb = MACFOR_CINZA
    run_placeholder.font.size = Pt(9)
    run_placeholder.font.italic = True

    run2 = paragraph.add_run()
    run2._r.append(fld_char_end)

    doc.add_page_break()


def _adicionar_tabela_metricas(doc, titulo, dados_linhas):
    """Adiciona uma tabela de métricas elegante.

    dados_linhas: list de tuples (metrica, atual, variacao_mes, variacao_ano)
    """
    doc.add_heading(titulo, level=2)

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Cabeçalho
    headers = ["Métrica", "Valor Atual", "Var. MoM", "Var. YoY"]
    for i, header_text in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header_text)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = MACFOR_BRANCO
        run.font.name = 'Calibri'
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COR_FUNDO_HEADER_TAB}"/>')
        cell._tc.get_or_add_tcPr().append(shading)

    # Linhas de dados
    for idx, (metrica, valor, var_mes, var_ano) in enumerate(dados_linhas):
        row = table.add_row()
        valores = [metrica, valor, var_mes, var_ano]
        for i, val in enumerate(valores):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            run.font.name = 'Calibri'
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i > 0 else WD_ALIGN_PARAGRAPH.LEFT

            # Colorir variações
            if i >= 2 and isinstance(val, str):
                if val.startswith('+'):
                    run.font.color.rgb = MACFOR_VERDE
                elif val.startswith('-'):
                    run.font.color.rgb = RGBColor(0xE7, 0x4C, 0x3C)

            # Fundo alternado
            if idx % 2 == 0:
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{COR_FUNDO_LINHA_ALT}"/>')
                cell._tc.get_or_add_tcPr().append(shading)

    # Larguras das colunas
    for row in table.rows:
        row.cells[0].width = Cm(5.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(3)
        row.cells[3].width = Cm(3)

    doc.add_paragraph()  # Espaçamento


def _formatar_variacao(valor):
    """Formata variação com sinal."""
    if valor > 0:
        return f"+{valor:.1f}%"
    elif valor < 0:
        return f"{valor:.1f}%"
    return "0.0%"


def _formatar_numero(valor, prefixo="", sufixo=""):
    """Formata número com separador de milhar."""
    if isinstance(valor, float):
        if abs(valor) < 100:
            return f"{prefixo}{valor:,.2f}{sufixo}"
        return f"{prefixo}{valor:,.0f}{sufixo}"
    return f"{prefixo}{valor:,}{sufixo}"


def _markdown_para_docx(doc, texto_md, nivel_base=2):
    """Converte texto Markdown simplificado em parágrafos do docx."""
    if not texto_md:
        return

    linhas = texto_md.split('\n')
    for linha in linhas:
        linha_strip = linha.strip()
        if not linha_strip:
            continue

        # Headings
        if linha_strip.startswith('### '):
            doc.add_heading(linha_strip[4:].strip().strip('*'), level=min(nivel_base + 1, 3))
        elif linha_strip.startswith('## '):
            doc.add_heading(linha_strip[3:].strip().strip('*'), level=nivel_base)
        elif linha_strip.startswith('# '):
            doc.add_heading(linha_strip[2:].strip().strip('*'), level=max(nivel_base - 1, 1))
        elif linha_strip.startswith('---'):
            # Linha horizontal — adiciona espaçamento sutil
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
        elif linha_strip.startswith(('- ', '* ', '• ')):
            # Lista com bullet
            texto_item = linha_strip.lstrip('-*• ').strip()
            p = doc.add_paragraph(style='List Bullet')
            _aplicar_formatacao_inline(p, texto_item)
        elif re.match(r'^\d+[\.\)] ', linha_strip):
            # Lista numerada
            texto_item = re.sub(r'^\d+[\.\)] ', '', linha_strip).strip()
            p = doc.add_paragraph(style='List Number')
            _aplicar_formatacao_inline(p, texto_item)
        else:
            p = doc.add_paragraph()
            _aplicar_formatacao_inline(p, linha_strip)


def _aplicar_formatacao_inline(paragraph, texto):
    """Aplica bold e italic inline no texto."""
    # Limpa o texto padrão do parágrafo
    paragraph.clear()

    # Regex para **bold** e *italic*
    partes = re.split(r'(\*\*.*?\*\*|\*.*?\*)', texto)
    for parte in partes:
        if parte.startswith('**') and parte.endswith('**'):
            run = paragraph.add_run(parte[2:-2])
            run.bold = True
        elif parte.startswith('*') and parte.endswith('*'):
            run = paragraph.add_run(parte[1:-1])
            run.italic = True
        else:
            paragraph.add_run(parte)


def gerar_docx_relatorio(dados, dados_investimentos, dados_custos, dados_seo,
                          etapa_cenario_atual, etapa_destaques, etapa_midias_pagas,
                          etapa_social, etapa_seo, etapa_aprendizados, etapa_proximos_passos):
    """Gera o relatório executivo completo em DOCX (pipeline de 7 etapas)."""

    doc = Document()
    mes_ref = datetime.now().strftime("%B/%Y").replace(
        "January", "Janeiro").replace("February", "Fevereiro").replace(
        "March", "Março").replace("April", "Abril").replace(
        "May", "Maio").replace("June", "Junho").replace(
        "July", "Julho").replace("August", "Agosto").replace(
        "September", "Setembro").replace("October", "Outubro").replace(
        "November", "Novembro").replace("December", "Dezembro")

    _configurar_estilos(doc)
    _adicionar_capa(doc, mes_ref)
    _adicionar_header_footer(doc, mes_ref)
    _adicionar_sumario(doc)

    # =================================================================
    # 1. COMPARATIVOS DE PERFORMANCE (tabelas numéricas)
    # =================================================================
    doc.add_heading("Comparativos de Performance", level=1)

    p = doc.add_paragraph()
    run = p.add_run("Visão consolidada dos principais indicadores de performance digital, "
                     "com variações mês a mês (MoM) e ano a ano (YoY).")
    run.font.italic = True
    run.font.color.rgb = MACFOR_CINZA

    metricas_perf = [
        ("Investimento", _formatar_numero(dados.get('spend_atual', 0), "R$ "),
         _formatar_variacao(dados.get('var_invest_mes', 0)),
         _formatar_variacao(dados.get('var_invest_ano', 0))),
        ("Sessões", _formatar_numero(dados.get('sess_atual', 0)),
         _formatar_variacao(dados.get('var_sess_mes', 0)),
         _formatar_variacao(dados.get('var_sess_ano', 0))),
        ("Alcance (Reach)", _formatar_numero(dados.get('reach_atual', 0)),
         _formatar_variacao(dados.get('var_reach_mes', 0)),
         _formatar_variacao(dados.get('var_reach_ano', 0))),
        ("Video Thruplays", _formatar_numero(dados.get('vtp_atual', 0)),
         _formatar_variacao(dados.get('var_vtp_mes', 0)),
         _formatar_variacao(dados.get('var_vtp_ano', 0))),
        ("Impressões", _formatar_numero(dados.get('imp_atual', 0)),
         _formatar_variacao(dados.get('var_imp_mes', 0)),
         _formatar_variacao(dados.get('var_imp_ano', 0))),
        ("Cliques", _formatar_numero(dados.get('cli_atual', 0)),
         _formatar_variacao(dados.get('var_cli_mes', 0)),
         _formatar_variacao(dados.get('var_cli_ano', 0))),
        ("Engajamentos", _formatar_numero(dados.get('eng_atual', 0)),
         _formatar_variacao(dados.get('var_eng_mes', 0)),
         _formatar_variacao(dados.get('var_eng_ano', 0))),
        ("CTR", _formatar_numero(dados.get('ctr_atual', 0), sufixo="%"),
         _formatar_variacao(dados.get('var_ctr_mes', 0)),
         _formatar_variacao(dados.get('var_ctr_ano', 0))),
    ]

    _adicionar_tabela_metricas(doc, "Indicadores Gerais", metricas_perf)

    inv_linhas = [
        ("Facebook", _formatar_numero(dados_investimentos.get('fb_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_fb_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_fb_ano', 0))),
        ("Instagram", _formatar_numero(dados_investimentos.get('ig_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_ig_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_ig_ano', 0))),
        ("TikTok", _formatar_numero(dados_investimentos.get('tt_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_tt_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_tt_ano', 0))),
        ("Google Ads", _formatar_numero(dados_investimentos.get('google_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_google_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_google_ano', 0))),
        ("Total", _formatar_numero(dados_investimentos.get('total_atual', 0), "R$ "),
         _formatar_variacao(dados_investimentos.get('var_total_mes', 0)),
         _formatar_variacao(dados_investimentos.get('var_total_ano', 0))),
    ]

    _adicionar_tabela_metricas(doc, "Investimentos por Canal", inv_linhas)

    custos_linhas = [
        ("CPC", _formatar_numero(dados_custos.get('cpc_atual', 0), "R$ "),
         _formatar_variacao(dados_custos.get('var_cpc_mes', 0)),
         _formatar_variacao(dados_custos.get('var_cpc_ano', 0))),
        ("CPM", _formatar_numero(dados_custos.get('cpm_atual', 0), "R$ "),
         _formatar_variacao(dados_custos.get('var_cpm_mes', 0)),
         _formatar_variacao(dados_custos.get('var_cpm_ano', 0))),
        ("CPE", _formatar_numero(dados_custos.get('cpe_atual', 0), "R$ "),
         _formatar_variacao(dados_custos.get('var_cpe_mes', 0)),
         _formatar_variacao(dados_custos.get('var_cpe_ano', 0))),
        ("CPV", _formatar_numero(dados_custos.get('cpv_atual', 0), "R$ "),
         _formatar_variacao(dados_custos.get('var_cpv_mes', 0)),
         _formatar_variacao(dados_custos.get('var_cpv_ano', 0))),
    ]

    _adicionar_tabela_metricas(doc, "Indicadores de Custo e Eficiência", custos_linhas)

    seo_linhas = [
        ("Visualizações (Total)", _formatar_numero(dados_seo.get('vis_total_atual', 0)),
         _formatar_variacao(dados_seo.get('var_vis_total_mes', 0)),
         _formatar_variacao(dados_seo.get('var_vis_total_ano', 0))),
        ("Sessões Orgânicas", _formatar_numero(dados_seo.get('sess_org_atual', 0)),
         _formatar_variacao(dados_seo.get('var_sess_org_mes', 0)),
         _formatar_variacao(dados_seo.get('var_sess_org_ano', 0))),
        ("Visualizações Orgânicas", _formatar_numero(dados_seo.get('vis_org_atual', 0)),
         _formatar_variacao(dados_seo.get('var_vis_org_mes', 0)),
         _formatar_variacao(dados_seo.get('var_vis_org_ano', 0))),
    ]
    _adicionar_tabela_metricas(doc, "Indicadores Orgânicos", seo_linhas)

    doc.add_page_break()

    # =================================================================
    # 2. CENÁRIO ATUAL
    # =================================================================
    doc.add_heading("Cenário Atual", level=1)
    _markdown_para_docx(doc, etapa_cenario_atual)
    doc.add_page_break()

    # =================================================================
    # 3. DESTAQUES
    # =================================================================
    doc.add_heading("Destaques do Período", level=1)
    _markdown_para_docx(doc, etapa_destaques)
    doc.add_page_break()

    # =================================================================
    # 4. MÍDIAS PAGAS
    # =================================================================
    doc.add_heading("Mídias Pagas", level=1)
    _markdown_para_docx(doc, etapa_midias_pagas)
    doc.add_page_break()

    # =================================================================
    # 5. SOCIAL
    # =================================================================
    doc.add_heading("Social", level=1)
    _markdown_para_docx(doc, etapa_social)
    doc.add_page_break()

    # =================================================================
    # 6. SEO
    # =================================================================
    doc.add_heading("SEO", level=1)
    _markdown_para_docx(doc, etapa_seo)
    doc.add_page_break()

    # =================================================================
    # 7. APRENDIZADOS
    # =================================================================
    doc.add_heading("Aprendizados", level=1)
    _markdown_para_docx(doc, etapa_aprendizados)
    doc.add_page_break()

    # =================================================================
    # 8. PRÓXIMOS PASSOS
    # =================================================================
    doc.add_heading("Próximos Passos", level=1)
    _markdown_para_docx(doc, etapa_proximos_passos)

    # =================================================================
    # RODAPÉ FINAL
    # =================================================================
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(8)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"Relatório gerado por IA em {datetime.now().strftime('%d/%m/%Y %H:%M')}")
    run.font.size = Pt(8)
    run.font.italic = True
    run.font.color.rgb = MACFOR_CINZA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Macfor Inteligência Digital | macfor.com.br")
    run.font.size = Pt(8)
    run.font.color.rgb = MACFOR_AZUL

    # Salvar em buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


def gerar_docx_cliente(dados, dados_investimentos, dados_custos, dados_seo,
                        etapa_cenario_atual, etapa_destaques, etapa_midias_pagas,
                        etapa_social, etapa_seo, etapa_aprendizados, etapa_proximos_passos):
    """Gera DOCX do relatório para o cliente — narrativo, elegante, focado em valor."""
    doc = Document()
    mes_ref = datetime.now().strftime("%B/%Y").replace(
        "January", "Janeiro").replace("February", "Fevereiro").replace(
        "March", "Março").replace("April", "Abril").replace(
        "May", "Maio").replace("June", "Junho").replace(
        "July", "Julho").replace("August", "Agosto").replace(
        "September", "Setembro").replace("October", "Outubro").replace(
        "November", "Novembro").replace("December", "Dezembro")

    _configurar_estilos(doc)

    # Capa diferenciada para o cliente
    for _ in range(6):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 60)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(10)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    run = p.add_run("RELATÓRIO DE RESULTADOS")
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = MACFOR_AZUL

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    run = p.add_run("Performance Digital & Inteligência Estratégica")
    run.font.size = Pt(14)
    run.font.color.rgb = MACFOR_CINZA

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    run = p.add_run("_" * 60)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(10)

    for _ in range(3):
        doc.add_paragraph()

    for label, valor in [("PREPARADO PARA", "Syngenta"), ("POR", "Macfor Inteligência Digital"), ("PERÍODO", mes_ref)]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"{label}: ")
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = MACFOR_AZUL
        run = p.add_run(valor)
        run.font.size = Pt(10)
        run.font.color.rgb = MACFOR_CINZA

    doc.add_page_break()
    _adicionar_header_footer(doc, mes_ref)
    _adicionar_sumario(doc)

    # Tabelas de performance
    doc.add_heading("Painel de Performance", level=1)
    p = doc.add_paragraph()
    run = p.add_run("Resultados consolidados do período com variações históricas.")
    run.font.italic = True
    run.font.color.rgb = MACFOR_CINZA

    metricas_perf = [
        ("Investimento", _formatar_numero(dados.get('spend_atual', 0), "R$ "),
         _formatar_variacao(dados.get('var_invest_mes', 0)), _formatar_variacao(dados.get('var_invest_ano', 0))),
        ("Alcance", _formatar_numero(dados.get('reach_atual', 0)),
         _formatar_variacao(dados.get('var_reach_mes', 0)), _formatar_variacao(dados.get('var_reach_ano', 0))),
        ("Impressões", _formatar_numero(dados.get('imp_atual', 0)),
         _formatar_variacao(dados.get('var_imp_mes', 0)), _formatar_variacao(dados.get('var_imp_ano', 0))),
        ("Cliques", _formatar_numero(dados.get('cli_atual', 0)),
         _formatar_variacao(dados.get('var_cli_mes', 0)), _formatar_variacao(dados.get('var_cli_ano', 0))),
        ("Engajamentos", _formatar_numero(dados.get('eng_atual', 0)),
         _formatar_variacao(dados.get('var_eng_mes', 0)), _formatar_variacao(dados.get('var_eng_ano', 0))),
        ("CTR", _formatar_numero(dados.get('ctr_atual', 0), sufixo="%"),
         _formatar_variacao(dados.get('var_ctr_mes', 0)), _formatar_variacao(dados.get('var_ctr_ano', 0))),
    ]
    _adicionar_tabela_metricas(doc, "Indicadores de Performance", metricas_perf)
    doc.add_page_break()

    # Cenário Atual
    doc.add_heading("Cenário Atual", level=1)
    _markdown_para_docx(doc, etapa_cenario_atual)
    doc.add_page_break()

    # Destaques
    doc.add_heading("Destaques do Período", level=1)
    _markdown_para_docx(doc, etapa_destaques)
    doc.add_page_break()

    # Mídias Pagas
    doc.add_heading("Mídias Pagas", level=1)
    _markdown_para_docx(doc, etapa_midias_pagas)
    doc.add_page_break()

    # Social
    doc.add_heading("Social", level=1)
    _markdown_para_docx(doc, etapa_social)
    doc.add_page_break()

    # SEO
    doc.add_heading("SEO", level=1)
    _markdown_para_docx(doc, etapa_seo)
    doc.add_page_break()

    # Aprendizados
    doc.add_heading("Aprendizados", level=1)
    _markdown_para_docx(doc, etapa_aprendizados)
    doc.add_page_break()

    # Próximos Passos
    doc.add_heading("Próximos Passos", level=1)
    _markdown_para_docx(doc, etapa_proximos_passos)

    # Rodapé
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("_" * 50)
    run.font.color.rgb = MACFOR_AZUL_CLARO
    run.font.size = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Macfor Inteligência Digital | macfor.com.br")
    run.font.size = Pt(8)
    run.font.color.rgb = MACFOR_AZUL

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# Estado da sessão para os cliques (Sincroniza com o BigQuery)
if 'bq_cliques' not in st.session_state:
    st.session_state.bq_cliques = {"atual": 0, "mes_passado": 0, "ano_passado": 0}

if 'relatorio_gerado' not in st.session_state:
    st.session_state.relatorio_gerado = False

# Título do aplicativo
st.title("📊 Relatório Executivo - Inteligência de Negócio")
st.markdown("*Transformando dados brutos em inteligência de mercado para seu cliente*")
st.markdown("---")

# Seletor de modelo de IA para geração de texto
col_modelo1, col_modelo2 = st.columns([1, 3])
with col_modelo1:
    modelos_disponiveis = ["Gemini"]
    if cliente_anthropic:
        modelos_disponiveis.append("Claude (Anthropic)")
    modelo_escolhido = st.selectbox("Motor de IA para texto:", modelos_disponiveis)
with col_modelo2:
    if modelo_escolhido == "Claude (Anthropic)":
        st.info("Claude gera o texto analítico. Gemini continua descrevendo as imagens.")
    else:
        st.info("Gemini gera texto e descreve imagens.")

# Estado da sessão
if 'relatorio_gerado' not in st.session_state:
    st.session_state.relatorio_gerado = False
if 'descricoes_imagens' not in st.session_state:
    st.session_state.descricoes_imagens = []
if 'descricoes_imagens_mes_passado' not in st.session_state:
    st.session_state.descricoes_imagens_mes_passado = []
if 'dados_processados' not in st.session_state:
    st.session_state.dados_processados = {}

# Função para descrever imagem (sempre usa Gemini - visão)
def descrever_imagem(imagem):
    try:
        response = modelo_visao.generate_content([
            """Analise este criativo de marketing digital com foco em INTELIGÊNCIA DE NEGÓCIO.
Descreva em até 150 palavras:
1. ELEMENTOS VISUAIS: cores, composição, imagens, texto sobreposto
2. MENSAGEM CENTRAL: o que está sendo comunicado e para qual público
3. ESTRATÉGIA CRIATIVA: qual técnica de persuasão está sendo usada (urgência, autoridade, identificação, prova social, etc.)
4. POTENCIAL DE PERFORMANCE: elementos que podem impactar positiva ou negativamente cliques, engajamento e conversão
5. DIFERENCIAÇÃO: o que torna este criativo único vs. padrões do mercado""",
            imagem
        ])
        return response.text
    except Exception as e:
        return f"Erro ao descrever imagem: {str(e)}"

# Função para calcular variações
def calcular_variacao(atual, anterior):
    if anterior and anterior != 0:
        return ((atual - anterior) / anterior) * 100
    return 0

#GERAÇÃO DE CONTEXTO ATUAL
#ADICIONE A COMPARAÇÃO DO ANO PASSADO E FAZER TODAS ANALISES CONSIDERANDO AS METRICAS DE PERFOMANCE PARA ALCANÇAR OS INDICADORES DE PERFORMANCE KPI OK
#FAZER O COMPARATIVO ENTRE A CONCORRENCIA E SHARE DE BUSCA OK
#COLOCAR DO INSTAGRAM, O AUMENTO DE SEGUIRDOR, SHARE DE CRESCIMENTO, VOLUME DE PUBLICAÇÃO
#COLOCAR O CRESCIMENTO EM CADA PLATAFORMA E POST DESTAQUE
#COLOCAR SNETIMENTO DO PUBLICO E PRINCIPAIS TÓPICOS

def gerar_cenario_atual(dados_metrica_performance, dados_investimentos, dados_custos, info_concorrentes, modelo_escolhido="Gemini"):
    """ETAPA 1/7: Cenário Atual — panorama geral da operação digital."""
    prompt = f"""
Você é um especialista sênior em marketing digital. Escreva a seção CENÁRIO ATUAL do relatório executivo mensal da Syngenta. Esta é a PRIMEIRA de 7 etapas e serve como base para todas as análises seguintes.

Escreva em prosa corrida, técnica e narrativa. Não use listas de bullet points como estrutura principal. Cada frase deve conter um dado concreto.

---

**DADOS DE PERFORMANCE:**
| Métrica | Atual | Var. MoM | Var. YoY |
|---------|-------|----------|----------|
| Investimento | R$ {dados_metrica_performance.get('spend_atual', 0):,.2f} | {dados_metrica_performance.get('var_invest_mes', 0):+.1f}%% | {dados_metrica_performance.get('var_invest_ano', 0):+.1f}%% |
| Sessões | {dados_metrica_performance.get('sess_atual', 0):,} | {dados_metrica_performance.get('var_sess_mes', 0):+.1f}%% | {dados_metrica_performance.get('var_sess_ano', 0):+.1f}%% |
| Alcance | {dados_metrica_performance.get('reach_atual', 0):,} | {dados_metrica_performance.get('var_reach_mes', 0):+.1f}%% | {dados_metrica_performance.get('var_reach_ano', 0):+.1f}%% |
| Impressões | {dados_metrica_performance.get('imp_atual', 0):,} | {dados_metrica_performance.get('var_imp_mes', 0):+.1f}%% | {dados_metrica_performance.get('var_imp_ano', 0):+.1f}%% |
| Cliques | {dados_metrica_performance.get('cli_atual', 0):,} | {dados_metrica_performance.get('var_cli_mes', 0):+.1f}%% | {dados_metrica_performance.get('var_cli_ano', 0):+.1f}%% |
| Engajamentos | {dados_metrica_performance.get('eng_atual', 0):,} | {dados_metrica_performance.get('var_eng_mes', 0):+.1f}%% | {dados_metrica_performance.get('var_eng_ano', 0):+.1f}%% |
| CTR | {dados_metrica_performance.get('ctr_atual', 0):.2f}%% | {dados_metrica_performance.get('var_ctr_mes', 0):+.1f}%% | {dados_metrica_performance.get('var_ctr_ano', 0):+.1f}%% |
| Video Thruplays | {dados_metrica_performance.get('vtp_atual', 0):,} | {dados_metrica_performance.get('var_vtp_mes', 0):+.1f}%% | {dados_metrica_performance.get('var_vtp_ano', 0):+.1f}%% |

**INVESTIMENTOS:**
Total: R$ {dados_investimentos.get('total_atual', 0):,.2f} (MoM: {dados_investimentos.get('var_total_mes', 0):+.1f}%%, YoY: {dados_investimentos.get('var_total_ano', 0):+.1f}%%)
Meta: R$ {dados_investimentos.get('fb_atual', 0) + dados_investimentos.get('ig_atual', 0):,.2f} | Google: R$ {dados_investimentos.get('google_atual', 0):,.2f} | TikTok: R$ {dados_investimentos.get('tt_atual', 0):,.2f}

**CUSTOS:**
CPC R$ {dados_custos.get('cpc_atual', 0):.2f} (MoM: {dados_custos.get('var_cpc_mes', 0):+.1f}%%, YoY: {dados_custos.get('var_cpc_ano', 0):+.1f}%%) | CPM R$ {dados_custos.get('cpm_atual', 0):.2f} (MoM: {dados_custos.get('var_cpm_mes', 0):+.1f}%%) | CPE R$ {dados_custos.get('cpe_atual', 0):.2f} (MoM: {dados_custos.get('var_cpe_mes', 0):+.1f}%%) | CPV R$ {dados_custos.get('cpv_atual', 0):.2f} (MoM: {dados_custos.get('var_cpv_mes', 0):+.1f}%%)

**CONCORRENTES:** {info_concorrentes if info_concorrentes else "Nenhuma informação fornecida."}

---

Cubra em sequência narrativa:

1. **Estado geral da operação**: crescimento, estabilidade ou retração? Classifique como EXPANSÃO, OTIMIZAÇÃO, ESTABILIDADE, CONTRAÇÃO ou CRISE com base nos 3 dados mais relevantes.

2. **Cruzamentos de dados obrigatórios** — identifique APENAS os que estão acontecendo nos dados:
- Investimento vs. Resultados: Efeito Tesoura (mais resultado com menos gasto) ou inverso?
- Impressões vs. CTR: relevância subindo ou caindo?
- Alcance vs. Sessões: público mais qualificado ou desconexão mídia-destino?
- CPC vs. CTR: relevância do anúncio melhorando ou piorando?
- Impressões vs. Alcance: frequência aumentando (risco de fadiga) ou alcance novo?

3. **Eficiência do funil**: calcule Impressões→Cliques (CTR), Cliques→Sessões (taxa de aterrissagem), Alcance→Engajamento. Onde o funil está apertando?

4. **Contexto competitivo**: pressão nos custos de mídia, sazonalidade agro, SWOT baseado exclusivamente nos dados.

5. **Veredicto executivo**: em 2-3 frases, sintetize o que aconteceu, por que, e o que significa.

**REGRAS:** Não invente dados. Tom técnico, narrativo, de especialista sênior. Português do Brasil.
"""
    return gerar_texto(prompt, modelo_escolhido)


def gerar_destaques(cenario_atual, modelo_escolhido="Gemini"):
    """ETAPA 2/7: Destaques do período."""
    prompt = f"""
Você é um especialista sênior em marketing digital. Escreva a seção DESTAQUES do relatório executivo mensal da Syngenta. Esta é a SEGUNDA de 7 etapas.

**CENÁRIO ATUAL (Etapa 1):**
{cenario_atual}

---

Extraia 5 a 7 fatos mais relevantes do período. Para cada destaque:

- **Título curto e impactante** com dado (ex: "Efeito Tesoura Confirmado: -12%% investimento, +8%% cliques")
- **Parágrafo narrativo denso**: o dado exato, contexto MoM/YoY, causa provável baseada nos cruzamentos da Etapa 1, e implicação prática.
- **Classificação**: CONQUISTA, ALERTA ou OPORTUNIDADE.

Cubra obrigatoriamente:
1. Principal ganho de eficiência — quantifique em economia ou produtividade
2. Métrica de maior crescimento — é sustentável ou pontual?
3. Principal sinal de alerta — cenário projetado se continuar
4. Correlação mais significativa entre métricas
5. Oportunidade concreta com maior potencial
6. Tendência geral da operação
7. (Opcional) Fato contraintuitivo

**REGRAS:** Não invente dados. Prosa corrida. Tom técnico. Português do Brasil.
"""
    return gerar_texto(prompt, modelo_escolhido)


def gerar_midias_pagas(cenario_atual, dados_investimentos, dados_custos, modelo_escolhido="Gemini"):
    """ETAPA 3/7: Análise de Mídias Pagas por canal."""
    prompt = f"""
Você é um especialista sênior em marketing digital. Escreva a seção MÍDIAS PAGAS do relatório executivo mensal da Syngenta. Esta é a TERCEIRA de 7 etapas.

**CENÁRIO ATUAL (Etapa 1):**
{cenario_atual}

---

**INVESTIMENTOS POR CANAL:**
| Canal | Atual | Var. MoM | Var. YoY |
|-------|-------|----------|----------|
| Facebook | R$ {dados_investimentos.get('fb_atual', 0):,.2f} | {dados_investimentos.get('var_fb_mes', 0):+.1f}%% | {dados_investimentos.get('var_fb_ano', 0):+.1f}%% |
| Instagram | R$ {dados_investimentos.get('ig_atual', 0):,.2f} | {dados_investimentos.get('var_ig_mes', 0):+.1f}%% | {dados_investimentos.get('var_ig_ano', 0):+.1f}%% |
| TikTok | R$ {dados_investimentos.get('tt_atual', 0):,.2f} | {dados_investimentos.get('var_tt_mes', 0):+.1f}%% | {dados_investimentos.get('var_tt_ano', 0):+.1f}%% |
| Google Ads | R$ {dados_investimentos.get('google_atual', 0):,.2f} | {dados_investimentos.get('var_google_mes', 0):+.1f}%% | {dados_investimentos.get('var_google_ano', 0):+.1f}%% |
| YouTube | R$ {dados_investimentos.get('yt_atual', 0):,.2f} | — | — |
| PMax | R$ {dados_investimentos.get('pmax_atual', 0):,.2f} | — | — |
| TOTAL | R$ {dados_investimentos.get('total_atual', 0):,.2f} | {dados_investimentos.get('var_total_mes', 0):+.1f}%% | {dados_investimentos.get('var_total_ano', 0):+.1f}%% |

**CUSTOS GLOBAIS:**
CPC R$ {dados_custos.get('cpc_atual', 0):.2f} (MoM: {dados_custos.get('var_cpc_mes', 0):+.1f}%%) | CPM R$ {dados_custos.get('cpm_atual', 0):.2f} (MoM: {dados_custos.get('var_cpm_mes', 0):+.1f}%%) | CPE R$ {dados_custos.get('cpe_atual', 0):.2f} (MoM: {dados_custos.get('var_cpe_mes', 0):+.1f}%%) | CPV R$ {dados_custos.get('cpv_atual', 0):.2f} (MoM: {dados_custos.get('var_cpv_mes', 0):+.1f}%%)

---

Analise em profundidade:

1. **Eficiência de capital por canal**: para cada canal, calcule %% do investimento total e cruze com participação nos resultados. Identifique o canal mais e menos eficiente.

2. **Análise por ecossistema**: Meta (FB+IG sinergia, audiência complementar ou sobreposta?), Google (Search+PMax, canibalização do orgânico?), TikTok (awareness real ou vanity?), YouTube (CPV competitivo?).

3. **Mix de mídia**: concentração de investimento (>50%% em um canal = risco). Proponha mix ideal baseado nos dados reais, não em teoria.

4. **Realocação sugerida**: se R$ X fossem movidos do canal A para o B, qual impacto estimado? Quantifique.

5. **Cruzamentos entre canais**: qual canal puxa eficiência para cima? Qual está inflacionando custos? Variações MoM alinhadas com resultados gerais?

**REGRAS:** Não invente dados. Prosa corrida. Tom técnico. Português do Brasil.
"""
    return gerar_texto(prompt, modelo_escolhido)


def gerar_social(cenario_atual, descricoes_imagens, descricoes_imagens_mes_passado, dados_custos, resumos_social_csvs=None, modelo_escolhido="Gemini"):
    """ETAPA 4/7: Análise de Social/Criativos."""

    bloco_csvs = ""
    bloco_analise_csv = ""
    num_recomendacoes = "6"
    if resumos_social_csvs:
        dados_csvs_juntos = chr(10).join(resumos_social_csvs)
        bloco_csvs = (
            "\n---\n\n"
            "**DADOS BRUTOS DE SOCIAL (CSVs enviados pelo usuário):**\n"
            "Os dados abaixo foram exportados diretamente das plataformas sociais. O formato varia — analise cada arquivo individualmente, "
            "identifique quais métricas e dimensões estão presentes, e extraia todos os insights possíveis. Cruze esses dados com os criativos, "
            "indicadores de eficiência e cenário atual.\n\n"
            f"{dados_csvs_juntos}\n\n---\n"
        )
        bloco_analise_csv = (
            "\n6. **Análise dos dados de Social (CSVs)**: analise todos os dados brutos fornecidos. "
            "Identifique padrões de engajamento, melhores horários, tipos de conteúdo com melhor performance, "
            "crescimento de seguidores, alcance orgânico vs. pago, e qualquer insight relevante que os dados revelem. "
            "Cruze com os criativos e indicadores de eficiência.\n\n"
            "7. **Cruzamentos entre fontes**: correlacione métricas dos CSVs com os dados de custo (CPE, CPC) "
            "e as análises visuais dos criativos. Identifique quais posts/conteúdos geraram melhor retorno.\n"
        )
        num_recomendacoes = "8"

    prompt = f"""
Você é um especialista sênior em marketing digital. Escreva a seção SOCIAL do relatório executivo mensal da Syngenta. Esta é a QUARTA de 7 etapas. Foque na análise de criativos, conteúdo social e performance criativa.

**CENÁRIO ATUAL (Etapa 1):**
{cenario_atual}

---

**CRIATIVOS DO MÊS ATUAL:**
{chr(10).join(descricoes_imagens) if descricoes_imagens else "Nenhum criativo do mês atual fornecido."}

**CRIATIVOS DO MÊS PASSADO:**
{chr(10).join(descricoes_imagens_mes_passado) if descricoes_imagens_mes_passado else "Nenhum criativo do mês passado fornecido."}

**INDICADORES DE EFICIÊNCIA CRIATIVA:**
| Indicador | Valor | Var. MoM | Var. YoY |
|-----------|-------|----------|----------|
| CPE | R$ {dados_custos.get('cpe_atual', 0):.2f} | {dados_custos.get('var_cpe_mes', 0):+.1f}%% | {dados_custos.get('var_cpe_ano', 0):+.1f}%% |
| CPC | R$ {dados_custos.get('cpc_atual', 0):.2f} | {dados_custos.get('var_cpc_mes', 0):+.1f}%% | {dados_custos.get('var_cpc_ano', 0):+.1f}%% |
| CPV | R$ {dados_custos.get('cpv_atual', 0):.2f} | {dados_custos.get('var_cpv_mes', 0):+.1f}%% | {dados_custos.get('var_cpv_ano', 0):+.1f}%% |
| CPM | R$ {dados_custos.get('cpm_atual', 0):.2f} | {dados_custos.get('var_cpm_mes', 0):+.1f}%% | {dados_custos.get('var_cpm_ano', 0):+.1f}%% |
{bloco_csvs}
---

Analise em profundidade:

1. **Estratégia narrativa**: qual mensagem central cada criativo comunica? Para qual público? Alinhamento com momento agro (safra, entressafra)?

2. **Psicologia aplicada**: quais gatilhos estão sendo usados (urgência, autoridade técnica, prova social, identificação com produtor rural)? Linguagem do campo ou genérica?

3. **Evolução mês a mês**: o que mudou? Mudança intencional ou aleatória? Elementos visuais mantidos como âncora de marca?

4. **ROI criativo**: correlacione mudanças visuais/narrativas com variações de CPE/CPC. Se CPE caiu, qual elemento causou? Se CPC subiu, o criativo gera curiosidade sem entregar promessa?

5. **Fadiga criativa**: criativos muito similares = saturação. Muito diferentes = inconsistência de marca.
{bloco_analise_csv}
{num_recomendacoes}. **Recomendações concretas**: testes A/B específicos, formatos a explorar (carrossel, Reels, UGC de produtor), ajustes de composição visual.

**REGRAS:** Não invente dados. Prosa corrida. Tom técnico. Português do Brasil. Se dados de CSVs foram fornecidos, OBRIGATORIAMENTE analise-os e extraia insights — não os ignore.
"""
    return gerar_texto(prompt, modelo_escolhido)


def gerar_seo_content(cenario_atual, dados_seo, dados_custos, resumos_seo_csvs=None, modelo_escolhido="Gemini"):
    """ETAPA 5/7: Análise de SEO e Conteúdo."""

    bloco_csvs = ""
    bloco_analise_seo_csv = ""
    if resumos_seo_csvs:
        dados_csvs_juntos = chr(10).join(resumos_seo_csvs)
        bloco_csvs = (
            "\n---\n\n"
            "**DADOS BRUTOS DE SEO (CSVs enviados pelo usuário):**\n"
            "Os dados abaixo foram exportados diretamente de ferramentas de SEO (Google Search Console, GA4, SEMrush, Ahrefs, etc.). "
            "O formato varia — analise cada arquivo individualmente, identifique quais métricas e dimensões estão presentes, "
            "e extraia todos os insights possíveis. Cruze esses dados com as métricas manuais acima e o cenário atual.\n\n"
            f"{dados_csvs_juntos}\n\n---\n"
        )
        bloco_analise_seo_csv = (
            "\n7. **Análise dos dados de SEO (CSVs)**: analise todos os dados brutos fornecidos. "
            "Identifique tendências de posicionamento, queries com potencial de crescimento, páginas de melhor/pior performance, "
            "CTR por posição, impressões vs. cliques, e qualquer insight relevante que os dados revelem.\n\n"
            "8. **Cruzamentos entre fontes**: correlacione os dados dos CSVs com as métricas manuais e o cenário atual. "
            "Identifique discrepâncias, oportunidades ocultas e confirme ou refute tendências.\n"
        )

    prompt = f"""
Você é um especialista sênior em marketing digital. Escreva a seção SEO do relatório executivo mensal da Syngenta. Esta é a QUINTA de 7 etapas.

**CENÁRIO ATUAL (Etapa 1):**
{cenario_atual}

---

**DADOS SEO:**
| Métrica | Atual | Mês Passado | Var. MoM |
|---------|-------|-------------|----------|
| Visualizações (Total) | {dados_seo.get('vis_total_atual', 0):,} | {dados_seo.get('vis_total_mes', 0):,} | {dados_seo.get('var_vis_total_mes', 0):+.1f}%% |
| Sessões (Total) | {dados_seo.get('sess_total_atual', 0):,} | {dados_seo.get('sess_total_mes', 0):,} | — |
| Usuários (Total) | {dados_seo.get('user_total_atual', 0):,} | {dados_seo.get('user_total_mes', 0):,} | — |
| Visualizações Orgânicas | {dados_seo.get('vis_org_atual', 0):,} | {dados_seo.get('vis_org_mes', 0):,} | {dados_seo.get('var_vis_org_mes', 0):+.1f}%% |
| Sessões Orgânicas | {dados_seo.get('sess_org_atual', 0):,} | {dados_seo.get('sess_org_mes', 0):,} | {dados_seo.get('var_sess_org_mes', 0):+.1f}%% |
| Usuários Orgânicos | {dados_seo.get('user_org_atual', 0):,} | {dados_seo.get('user_org_mes', 0):,} | — |

**Top Keywords:** {dados_seo.get('top_keywords', 'Nenhuma keyword fornecida')}

**CPC Médio (para cálculo de custo evitado):** R$ {dados_custos.get('cpc_atual', 0):.2f}
{bloco_csvs}
---

Analise em profundidade:

1. **Demanda de mercado via buscas**: keywords de marca vs. genéricas? Proporção e implicação estratégica.

2. **Autoridade de marca**: orgânico crescendo = ganhando autoridade. Caindo = concorrentes produzindo conteúdo melhor ou conteúdo não responde às perguntas do público.

3. **Funil de conteúdo**: visualizações→sessões→usuários. Taxas de conversão entre etapas. Onde há gargalo?

4. **Cruzamento orgânico vs. pago**: proporção de tráfego orgânico sobre total. Independência de mídia crescendo ou dependência de pago aumentando?

5. **Custo evitado pelo orgânico**: sessões orgânicas × CPC médio = economia de mídia paga. Quantifique.

6. **Content gap analysis**: baseado nas keywords, quais temas o público busca que a Syngenta não cobre?
{bloco_analise_seo_csv}
**REGRAS:** Não invente dados. Prosa corrida. Tom técnico. Português do Brasil. Se dados de CSVs foram fornecidos, OBRIGATORIAMENTE analise-os e extraia insights — não os ignore.
"""
    return gerar_texto(prompt, modelo_escolhido)


def gerar_aprendizados(cenario_atual, destaques, midias_pagas, social, seo, dados_metrica_performance, dados_custos, dados_seo, modelo_escolhido="Gemini"):
    """ETAPA 6/7: Aprendizados — diagnóstico de eficiência, red flags, oportunidades."""
    prompt = f"""
Você é um especialista sênior em marketing digital. Escreva a seção APRENDIZADOS do relatório executivo mensal da Syngenta. Esta é a SEXTA de 7 etapas. Consolide tudo que foi analisado em insights acionáveis.

**ETAPAS ANTERIORES:**

Cenário Atual:
{cenario_atual}

Destaques:
{destaques}

Mídias Pagas:
{midias_pagas}

Social:
{social}

SEO:
{seo}

---

**DADOS COMPLEMENTARES:**
Cliques: {dados_metrica_performance.get('cli_atual', 0):,} | Engajamentos: {dados_metrica_performance.get('eng_atual', 0):,} | Sessões: {dados_metrica_performance.get('sess_atual', 0):,}
CPC: R$ {dados_custos.get('cpc_atual', 0):.2f} (MoM: {dados_custos.get('var_cpc_mes', 0):+.1f}%%) | CPM: R$ {dados_custos.get('cpm_atual', 0):.2f} (MoM: {dados_custos.get('var_cpm_mes', 0):+.1f}%%) | CPE: R$ {dados_custos.get('cpe_atual', 0):.2f} (MoM: {dados_custos.get('var_cpe_mes', 0):+.1f}%%) | CPV: R$ {dados_custos.get('cpv_atual', 0):.2f} (MoM: {dados_custos.get('var_cpv_mes', 0):+.1f}%%)
Orgânico: {dados_seo.get('vis_org_atual', 0):,} visualizações

---

Cubra em sequência narrativa:

1. **Diagnóstico de eficiência operacional**:
   - Score de Saúde Digital (1-10) com critérios explícitos (base 5, +2 se custos otimizando, +2 se Efeito Tesoura, +1 orgânico crescendo, +1 funil sem gargalos, e inversos)
   - Classifique cada custo (CPC/CPM/CPE/CPV) como OTIMIZANDO, ESTÁVEL ou INFLACIONANDO
   - Economia real calculada em reais para cada custo que caiu
   - Efeito Tesoura: confirme ou negue, quantifique

2. **Red flags e pontos de atenção**: para cada alerta identificado nas etapas anteriores, construa ficha com: SINAL, EVIDÊNCIA CRUZADA, CAUSA PROVÁVEL, IMPACTO QUANTIFICADO, AÇÃO RECOMENDADA, URGÊNCIA (CRÍTICA/ALTA/MÉDIA/BAIXA). Se operação saudável, documente por quê.

3. **Mapa de oportunidades**: 3-5 oportunidades ordenadas por impacto/esforço:
   - QUICK WINS (alto impacto, baixo esforço, até 7 dias)
   - MOVIMENTOS TÁTICOS (alto impacto, médio esforço, 2-4 semanas)
   - APOSTAS ESTRATÉGICAS (alto impacto, alto esforço, longo prazo)
   Para cada: evidência nos dados, potencial quantificado, ação específica.

**REGRAS:** Não invente dados. NÃO repita o que já foi dito — sintetize e avance. Prosa corrida. Tom honesto e técnico. Português do Brasil.
"""
    return gerar_texto(prompt, modelo_escolhido)


def gerar_proximos_passos(cenario_atual, aprendizados, modelo_escolhido="Gemini"):
    """ETAPA 7/7: Próximos Passos — plano de ação concreto."""
    prompt = f"""
Você é um especialista sênior em marketing digital. Escreva a seção PRÓXIMOS PASSOS do relatório executivo mensal da Syngenta. Esta é a SÉTIMA e ÚLTIMA etapa.

**CENÁRIO ATUAL (Etapa 1):**
{cenario_atual}

**APRENDIZADOS (Etapa 6):**
{aprendizados}

---

Construa um plano de ação concreto e executável:

1. **Inteligência acumulada**: 3-5 descobertas mais importantes do período. Para cada: o insight em uma frase, o dado que sustenta, a implicação para os próximos 30-90 dias.

2. **Ações imediatas (0-30 dias)**: para cada ação especifique o que fazer, área responsável (mídia, criativo, conteúdo, estratégia), prazo, KPI de sucesso. Prioridades: resolver red flags urgentes, implementar quick wins, ajustes de budget, testes A/B.

3. **Movimentos táticos (30-60 dias)**: para cada: objetivo, recurso necessário, resultado esperado com meta numérica. Prioridades: novos canais/formatos, reestruturação de campanhas, programa de conteúdo, otimização de landing pages.

4. **Visão estratégica (60-180 dias)**: para cada diretriz: a tese, evidências nos dados, investimento estimado, retorno projetado. Temas: mix de mídia ideal, meta de %% orgânico, ativos digitais proprietários, posicionamento competitivo.

5. **Matriz de priorização final**: ordene TODAS as ações em sequência lógica de execução, indicando dependências e conflitos de recurso.

**REGRAS:** Não invente dados. Máxima especificidade. Tom técnico e direto. Português do Brasil.
"""
    return gerar_texto(prompt, modelo_escolhido)


if st.button("🔄 Atualizar Dados (Syngenta)"):
    with st.spinner("Buscando dados históricos no BigQuery..."):
        res = fetch_bigquery_data()
    
        if res:
            # Investimento (Spend)
            st.session_state.spend_atual = float(res.get('spend_atual') or 0.0)
            st.session_state.spend_mes = float(res.get('spend_mes') or 0.0)
            st.session_state.spend_ano = float(res.get('spend_ano') or 0.0)

            # Sessões (Sess)
            st.session_state.sess_atual = int(res.get('sess_atual') or 0)
            st.session_state.sess_mes = int(res.get('sess_mes') or 0)
            st.session_state.sess_ano = int(res.get('sess_ano') or 0)

            # Alcance (Reach)
            st.session_state.reach_atual = int(res.get('reach_atual') or 0)
            st.session_state.reach_mes = int(res.get('reach_mes') or 0)
            st.session_state.reach_ano = int(res.get('reach_ano') or 0)

            # Video Thruplays (VTP)
            st.session_state.vtp_atual = int(res.get('vtp_atual') or 0)
            st.session_state.vtp_mes = int(res.get('vtp_mes') or 0)
            st.session_state.vtp_ano = int(res.get('vtp_ano') or 0)

            # Cliques
            st.session_state.cli_atual = int(res.get('cli_atual') or 0)
            st.session_state.cli_mes = int(res.get('cli_mes') or 0)
            st.session_state.cli_ano = int(res.get('cli_ano') or 0)
            
            # Impressões
            st.session_state.imp_atual = int(res.get('imp_atual') or 0)
            st.session_state.imp_mes = int(res.get('imp_mes') or 0)
            st.session_state.imp_ano = int(res.get('imp_ano') or 0)
            
            # Engajamentos
            st.session_state.eng_atual = int(res.get('eng_atual') or 0)
            st.session_state.eng_mes = int(res.get('eng_mes') or 0)
            st.session_state.eng_ano = int(res.get('eng_ano') or 0)

            
            # Atribuindo diretamente às chaves que os st.number_input usam como 'key'
            st.session_state.cpc_atual = float(res.get('cpc_atual') or 0.0)
            st.session_state.cpc_mes = float(res.get('cpc_mes') or 0.0)
            st.session_state.cpc_ano = float(res.get('cpc_ano') or 0.0)

            st.session_state.cpm_atual = float(res.get('cpm_atual') or 0.0)
            st.session_state.cpm_mes = float(res.get('cpm_mes') or 0.0)
            st.session_state.cpm_ano = float(res.get('cpm_ano') or 0.0)

            # CTR (conforme solicitado anteriormente)
            st.session_state.ctr_atual = float(res.get('ctr_atual') or 0.0)
            st.session_state.ctr_mes = float(res.get('ctr_mes') or 0.0)
            st.session_state.ctr_ano = float(res.get('ctr_ano') or 0.0)

            # Facebook
            st.session_state.fb_atual = float(res.get('spend_fb_atual') or 0.0)
            st.session_state.fb_mes = float(res.get('spend_fb_mes') or 0.0)
            st.session_state.fb_ano = float(res.get('spend_fb_ano') or 0.0)
            
            # Google Ads (Pode mapear para seu campo de Google Ads ou somar no campo de 'Display/YT/PMax' se preferir)
            st.session_state.google_atual = float(res.get('spend_google_atual') or 0.0)
            st.session_state.google_mes = float(res.get('spend_google_mes') or 0.0)
            st.session_state.google_ano = float(res.get('spend_google_ano') or 0.0)

            # TikTok
            st.session_state.tt_atual = float(res.get('spend_tiktok_atual') or 0.0)
            st.session_state.tt_mes = float(res.get('spend_tiktok_mes') or 0.0)
            st.session_state.tt_ano = float(res.get('spend_tiktok_ano') or 0.0)


            st.success("Dados sincronizados com sucesso!")
            st.rerun()

# Formulário principal
with st.form("relatorio_form"):
    st.header("📝 Dados do Relatório")
    
    # Contexto e Destaques
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Contexto Geral")
        contexto_input = st.text_area("Contexto Atual (opcional)", height=100, placeholder="Descreva o contexto da campanha/período...")
        info_concorrentes = st.text_area("Informações de Concorrentes", height=100, placeholder="O que os concorrentes estão fazendo?")
    
    with col2:
        st.subheader("Upload de Criativos")
        st.markdown("**Criativos do Mês Atual**")
        imagens = st.file_uploader("Faça upload dos criativos atuais", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True, key="upload_atual")
        st.markdown("**Criativos do Mês Passado** *(para comparação)*")
        imagens_mes_passado = st.file_uploader("Faça upload dos criativos do mês passado", type=['png', 'jpg', 'jpeg'], accept_multiple_files=True, key="upload_mes_passado")

    # Dados de Social (CSVs)
    st.subheader("📱 Dados de Social (CSVs)")
    st.markdown(
        "Faça upload de CSVs exportados das plataformas sociais (Instagram Insights, TikTok Analytics, "
        "Facebook Page, X/Twitter, LinkedIn, etc.). **Qualquer formato é aceito** — o sistema lê os dados "
        "brutos e cruza com as demais informações do relatório via IA generativa."
    )
    social_csvs = st.file_uploader(
        "Upload de CSVs de Social",
        type=['csv'],
        accept_multiple_files=True,
        key="upload_social_csvs",
        help="Aceita múltiplos arquivos. O formato pode variar — o sistema não assume nenhuma estrutura fixa."
    )

# Métricas Principais
    st.subheader("📊 Métricas de Performance")
    
    # Criamos os cabeçalhos das colunas
    col_label1, col_label2, col_label3 = st.columns(3)
    with col_label1: st.markdown("### **Atual**")
    with col_label2: st.markdown("### **Mês Passado**")
    with col_label3: st.markdown("### **Ano Passado**")

    # Lista de métricas para iterar e criar os campos dinamicamente ou manualmente
    # Vou organizar em blocos para manter a estrutura que você pediu
    
    def criar_linha_metrica(label, key_prefix):
        c1, c2, c3 = st.columns(3)
        with c1:
            st.number_input(f"{label}", min_value=0.0, key=f"{key_prefix}_atual", format="%.2f" if "invest" in key_prefix or "c" in key_prefix else "%.0f")
        with c2:
            st.number_input(f"{label}", min_value=0.0, key=f"{key_prefix}_mes", format="%.2f" if "invest" in key_prefix or "c" in key_prefix else "%.0f")
        with c3:
            st.number_input(f"{label}", min_value=0.0, key=f"{key_prefix}_ano", format="%.2f" if "invest" in key_prefix or "c" in key_prefix else "%.0f")

    # Renderizando os blocos solicitados
    st.markdown("---")
    criar_linha_metrica("Investimento", "spend")
    criar_linha_metrica("Sessões", "sess")
    criar_linha_metrica("Alcance (Reach)", "reach")
    criar_linha_metrica("Video Thruplays", "vtp")
    criar_linha_metrica("Visualizações", "vis")
    criar_linha_metrica("Impressões", "imp")
    criar_linha_metrica("Cliques", "cli")
    criar_linha_metrica("Engajamentos", "eng")
    criar_linha_metrica("CTR (%)", "ctr")

    # Investimentos
    st.subheader("💰 Investimentos por Canal")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Atual**")
        investimento_fb_atual = st.number_input("Facebook", min_value=0.0, format="%.2f", key="fb_atual")
        investimento_ig_atual = st.number_input("Instagram", min_value=0.0, format="%.2f", key="ig_atual")
        investimento_tt_atual = st.number_input("TikTok", min_value=0.0, format="%.2f", key="tt_atual")
    
    with col2:
        st.markdown("**Mês Passado**")
        investimento_fb_mes_passado = st.number_input("Facebook", min_value=0.0, format="%.2f", key="fb_mes")
        investimento_ig_mes_passado = st.number_input("Instagram", min_value=0.0, format="%.2f", key="ig_mes")
        investimento_tt_mes_passado = st.number_input("TikTok", min_value=0.0, format="%.2f", key="tt_mes")
    
    with col3:
        st.markdown("**Ano Passado**")
        investimento_fb_ano_passado = st.number_input("Facebook", min_value=0.0, format="%.2f", key="fb_ano")
        investimento_ig_ano_passado = st.number_input("Instagram", min_value=0.0, format="%.2f", key="ig_ano")
        investimento_tt_ano_passado = st.number_input("TikTok", min_value=0.0, format="%.2f", key="tt_ano")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        investimento_ads_atual = st.number_input("Google Ads (Atual)", min_value=0.0, format="%.2f", key="google_atual")
        investimento_yt_atual = st.number_input("YouTube (Atual)", min_value=0.0, format="%.2f")
        investimento_pmax_atual = st.number_input("PMax (Atual)", min_value=0.0, format="%.2f")
    
    with col2:
        investimento_ads_mes_passado = st.number_input("Google Ads (Mês Passado)", min_value=0.0, format="%.2f", key="google_mes")
        investimento_yt_mes_passado = st.number_input("YouTube (Mês Passado)", min_value=0.0, format="%.2f")
        investimento_pmax_mes_passado = st.number_input("PMax (Mês Passado)", min_value=0.0, format="%.2f")
    
    with col3:
        investimento_ads_ano_passado = st.number_input("Google Ads (Ano Passado)", min_value=0.0, format="%.2f", key="google_ano")
        investimento_yt_ano_passado = st.number_input("YouTube (Ano Passado)", min_value=0.0, format="%.2f")
        investimento_pmax_ano_passado = st.number_input("PMax (Ano Passado)", min_value=0.0, format="%.2f")
    
    # Custos
    st.subheader("💰 Custos")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Atual**")
        cpe_atual = st.number_input("Custo por Engajamento", min_value=0.0, format="%.2f", key="cpe_atual")
        cpc_atual = st.number_input("Custo por Clique", min_value=0.0, format="%.2f", key="cpc_atual")
        cpv_atual = st.number_input("Custo por Visualização", min_value=0.0, format="%.2f", key="cpv_atual")
        cpm_atual = st.number_input("Custo por Mil Impressões", min_value=0.0, format="%.2f", key="cpm_atual")
    
    with col2:
        st.markdown("**Mês Passado**")
        cpe_mes_passado = st.number_input("Custo por Engajamento", min_value=0.0, format="%.2f", key="cpe_mes")
        cpc_mes_passado = st.number_input("Custo por Clique", min_value=0.0, format="%.2f", key="cpc_mes")
        cpv_mes_passado = st.number_input("Custo por Visualização", min_value=0.0, format="%.2f", key="cpv_mes")
        cpm_mes_passado = st.number_input("Custo por Mil Impressões", min_value=0.0, format="%.2f", key="cpm_mes")
    
    with col3:
        st.markdown("**Ano Passado**")
        cpe_ano_passado = st.number_input("Custo por Engajamento", min_value=0.0, format="%.2f", key="cpe_ano")
        cpc_ano_passado = st.number_input("Custo por Clique", min_value=0.0, format="%.2f", key="cpc_ano")
        cpv_ano_passado = st.number_input("Custo por Visualização", min_value=0.0, format="%.2f", key="cpv_ano")
        cpm_ano_passado = st.number_input("Custo por Mil Impressões", min_value=0.0, format="%.2f", key="cpm_ano")
    
    # SEO + Content
    st.subheader("🔍 SEO + Content")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        st.markdown("**Atual**")
        seo_visualizacoes_atual = st.number_input("Visualizações", min_value=0, key="seo_vis_atual")
        seo_sessoes_atual = st.number_input("Sessões", min_value=0, key="seo_sess_atual")
        seo_usuarios_atual = st.number_input("Usuários", min_value=0, key="seo_user_atual")
    
    with col2:
        st.markdown("**Mês Passado**")
        seo_visualizacoes_mes_passado = st.number_input("Visualizações", min_value=0, key="seo_vis_mes")
        seo_sessoes_mes_passado = st.number_input("Sessões", min_value=0, key="seo_sess_mes")
        seo_usuarios_mes_passado = st.number_input("Usuários", min_value=0, key="seo_user_mes")
    
    with col3:
        st.markdown("**Ano Passado**")
        seo_visualizacoes_ano_passado = st.number_input("Visualizações", min_value=0, key="seo_vis_ano")
        seo_sessoes_ano_passado = st.number_input("Sessões", min_value=0, key="seo_sess_ano")
        seo_usuarios_ano_passado = st.number_input("Usuários", min_value=0, key="seo_user_ano")
    
    col1, col2, col3 = st.columns(3)
    with col1:
        seo_visualizacoes_org_atual = st.number_input("Visualizações Orgânicas (Atual)", min_value=0)
        seo_sessoes_org_atual = st.number_input("Sessões Orgânicas (Atual)", min_value=0)
        seo_usuarios_org_atual = st.number_input("Usuários Orgânicos (Atual)", min_value=0)
    
    with col2:
        seo_visualizacoes_org_mes_passado = st.number_input("Visualizações Orgânicas (Mês Passado)", min_value=0)
        seo_sessoes_org_mes_passado = st.number_input("Sessões Orgânicas (Mês Passado)", min_value=0)
        seo_usuarios_org_mes_passado = st.number_input("Usuários Orgânicos (Mês Passado)", min_value=0)
    
    with col3:
        seo_visualizacoes_org_ano_passado = st.number_input("Visualizações Orgânicas (Ano Passado)", min_value=0)
        seo_sessoes_org_ano_passado = st.number_input("Sessões Orgânicas (Ano Passado)", min_value=0)
        seo_usuarios_org_ano_passado = st.number_input("Usuários Orgânicos (Ano Passado)", min_value=0)
    
    top_keywords = st.text_area("Top 10 Palavras-chave do Mês", height=100, placeholder="Liste as principais palavras-chave...")

    st.markdown("**📂 Dados de SEO (CSVs)**")
    st.markdown(
        "Faça upload de CSVs exportados de ferramentas de SEO (Google Search Console, GA4, SEMrush, Ahrefs, etc.). "
        "**Qualquer formato é aceito** — o sistema lê os dados brutos e cruza com as demais informações via IA."
    )
    seo_csvs = st.file_uploader(
        "Upload de CSVs de SEO",
        type=['csv'],
        accept_multiple_files=True,
        key="upload_seo_csvs",
        help="Aceita múltiplos arquivos. O formato pode variar — o sistema não assume nenhuma estrutura fixa."
    )
    
    submitted = st.form_submit_button("🚀 Gerar Relatório Executivo")

if submitted:
    # Calcular totais e variações
    investimento_total_atual = (investimento_fb_atual + investimento_ig_atual + investimento_tt_atual + 
                                investimento_ads_atual + investimento_yt_atual + investimento_pmax_atual)
    investimento_total_mes_passado = (investimento_fb_mes_passado + investimento_ig_mes_passado + investimento_tt_mes_passado + 
                                      investimento_ads_mes_passado + investimento_yt_mes_passado + investimento_pmax_mes_passado)
    investimento_total_ano_passado = (investimento_fb_ano_passado + investimento_ig_ano_passado + investimento_tt_ano_passado + 
                                      investimento_ads_ano_passado + investimento_yt_ano_passado + investimento_pmax_ano_passado)
    
    # Processar imagens do mês atual (sempre com Gemini)
    descricoes_imagens = []
    if imagens:
        with st.spinner("Analisando criativos do mês atual..."):
            for imagem_file in imagens:
                image = Image.open(imagem_file)
                descricao = descrever_imagem(image)
                descricoes_imagens.append(f"**[ATUAL] {imagem_file.name}**: {descricao}")

    # Processar imagens do mês passado (sempre com Gemini)
    descricoes_imagens_mes_passado = []
    if imagens_mes_passado:
        with st.spinner("Analisando criativos do mês passado..."):
            for imagem_file in imagens_mes_passado:
                image = Image.open(imagem_file)
                descricao = descrever_imagem(image)
                descricoes_imagens_mes_passado.append(f"**[MES PASSADO] {imagem_file.name}**: {descricao}")

    # Processar CSVs de Social — leitura agnóstica de formato
    resumos_social_csvs = []
    if social_csvs:
        with st.spinner("Lendo dados de Social (CSVs)..."):
            for csv_file in social_csvs:
                try:
                    df = pd.read_csv(csv_file)
                    # Limitar preview para não estourar contexto do LLM
                    max_linhas = 80
                    preview = df.head(max_linhas).to_markdown(index=False)
                    resumo_estatistico = df.describe(include='all').to_markdown()
                    resumo = (
                        f"### Arquivo: {csv_file.name}\n"
                        f"- **Linhas:** {len(df)} | **Colunas:** {len(df.columns)}\n"
                        f"- **Colunas:** {', '.join(df.columns.tolist())}\n\n"
                        f"**Resumo estatístico:**\n{resumo_estatistico}\n\n"
                        f"**Dados ({min(len(df), max_linhas)} primeiras linhas):**\n{preview}"
                    )
                    resumos_social_csvs.append(resumo)
                except Exception as e:
                    resumos_social_csvs.append(f"### Arquivo: {csv_file.name}\n⚠️ Erro ao ler: {str(e)}")

    # Processar CSVs de SEO — leitura agnóstica de formato
    resumos_seo_csvs = []
    if seo_csvs:
        with st.spinner("Lendo dados de SEO (CSVs)..."):
            for csv_file in seo_csvs:
                try:
                    df = pd.read_csv(csv_file)
                    max_linhas = 80
                    preview = df.head(max_linhas).to_markdown(index=False)
                    resumo_estatistico = df.describe(include='all').to_markdown()
                    resumo = (
                        f"### Arquivo: {csv_file.name}\n"
                        f"- **Linhas:** {len(df)} | **Colunas:** {len(df.columns)}\n"
                        f"- **Colunas:** {', '.join(df.columns.tolist())}\n\n"
                        f"**Resumo estatístico:**\n{resumo_estatistico}\n\n"
                        f"**Dados ({min(len(df), max_linhas)} primeiras linhas):**\n{preview}"
                    )
                    resumos_seo_csvs.append(resumo)
                except Exception as e:
                    resumos_seo_csvs.append(f"### Arquivo: {csv_file.name}\n⚠️ Erro ao ler: {str(e)}")

    # Organizar dados
    dados_metrica_performance = {
        'spend_atual': st.session_state.get('spend_atual', 0),
        'sess_atual': st.session_state.get('sess_atual', 0),
        'reach_atual': st.session_state.get('reach_atual', 0),
        'vtp_atual': st.session_state.get('vtp_atual', 0),   
        'vis_atual': st.session_state.get('vis_atual', 0),   
        'imp_atual': st.session_state.get('imp_atual', 0),
        'cli_atual': st.session_state.get('cli_atual', 0),
        'eng_atual': st.session_state.get('eng_atual', 0),
        'ctr_atual': st.session_state.get('ctr_atual', 0),

        # --- VARIAÇÕES ANO SOBRE ANO (YoY - 2026 vs 2025) ---
        'var_invest_ano': calcular_variacao(st.session_state.get('spend_atual', 0), st.session_state.get('spend_ano', 0)),
        'var_sess_ano':   calcular_variacao(st.session_state.get('sess_atual', 0),  st.session_state.get('sess_ano', 0)),
        'var_reach_ano':  calcular_variacao(st.session_state.get('reach_atual', 0), st.session_state.get('reach_ano', 0)),
        'var_vtp_ano':    calcular_variacao(st.session_state.get('vtp_atual', 0),   st.session_state.get('vtp_ano', 0)),
        'var_vis_ano':    calcular_variacao(st.session_state.get('vis_atual', 0),   st.session_state.get('vis_ano', 0)),
        'var_imp_ano':    calcular_variacao(st.session_state.get('imp_atual', 0),   st.session_state.get('imp_ano', 0)),
        'var_cli_ano':    calcular_variacao(st.session_state.get('cli_atual', 0),   st.session_state.get('cli_ano', 0)),
        'var_eng_ano':    calcular_variacao(st.session_state.get('eng_atual', 0),   st.session_state.get('eng_ano', 0)),
        'var_ctr_ano':    calcular_variacao(st.session_state.get('ctr_atual', 0),   st.session_state.get('ctr_ano', 0)),

        # --- VARIAÇÕES MÊS SOBRE MÊS (MoM - Atual vs Mês Passado) ---
        'var_invest_mes': calcular_variacao(st.session_state.get('spend_atual', 0), st.session_state.get('spend_mes', 0)),
        'var_sess_mes':   calcular_variacao(st.session_state.get('sess_atual', 0),  st.session_state.get('sess_mes', 0)),
        'var_reach_mes':  calcular_variacao(st.session_state.get('reach_atual', 0), st.session_state.get('reach_mes', 0)),
        'var_vtp_mes':    calcular_variacao(st.session_state.get('vtp_atual', 0),   st.session_state.get('vtp_mes', 0)),
        'var_vis_mes':    calcular_variacao(st.session_state.get('vis_atual', 0),   st.session_state.get('vis_mes', 0)),
        'var_imp_mes':    calcular_variacao(st.session_state.get('imp_atual', 0),   st.session_state.get('imp_mes', 0)),
        'var_cli_mes':    calcular_variacao(st.session_state.get('cli_atual', 0),   st.session_state.get('cli_mes', 0)),
        'var_eng_mes':    calcular_variacao(st.session_state.get('eng_atual', 0),   st.session_state.get('eng_mes', 0)),
        'var_ctr_mes':    calcular_variacao(st.session_state.get('ctr_atual', 0),   st.session_state.get('ctr_mes', 0)),

        # --- DADOS EXTRAS PARA COMPOSIÇÃO ---
        'info_concorrentes': info_concorrentes,
        'contexto_input': contexto_input,
        'cpe_atual': st.session_state.get('cpe_atual', 0),
        'cpc_atual': st.session_state.get('cpc_atual', 0),
        
        
        # Informações adicionais
        'info_concorrentes': info_concorrentes,
        'top_keywords': top_keywords,
        'contexto_input': contexto_input
    }

    # DICIONÁRIO EXCLUSIVO DE CUSTOS E EFICIÊNCIA
    dados_custos = {
        # --- VALORES ATUAIS (O que o usuário vê na tela) ---
        'cpe_atual': st.session_state.get('cpe_atual', 0),
        'cpc_atual': st.session_state.get('cpc_atual', 0),
        'cpv_atual': st.session_state.get('cpv_atual', 0),
        'cpm_atual': st.session_state.get('cpm_atual', 0),

        # --- VARIAÇÕES MÊS SOBRE MÊS (MoM - Eficiência Mensal) ---
        'var_cpe_mes': calcular_variacao(st.session_state.get('cpe_atual', 0), st.session_state.get('cpe_mes', 0)),
        'var_cpc_mes': calcular_variacao(st.session_state.get('cpc_atual', 0), st.session_state.get('cpc_mes', 0)),
        'var_cpv_mes': calcular_variacao(st.session_state.get('cpv_atual', 0), st.session_state.get('cpv_mes', 0)),
        'var_cpm_mes': calcular_variacao(st.session_state.get('cpm_atual', 0), st.session_state.get('cpm_mes', 0)),

        # --- VARIAÇÕES ANO SOBRE ANO (YoY - Eficiência Histórica) ---
        'var_cpe_ano': calcular_variacao(st.session_state.get('cpe_atual', 0), st.session_state.get('cpe_ano', 0)),
        'var_cpc_ano': calcular_variacao(st.session_state.get('cpc_atual', 0), st.session_state.get('cpc_ano', 0)),
        'var_cpv_ano': calcular_variacao(st.session_state.get('cpv_atual', 0), st.session_state.get('cpv_ano', 0)),
        'var_cpm_ano': calcular_variacao(st.session_state.get('cpm_atual', 0), st.session_state.get('cpm_ano', 0)),
    }
    # DICIONÁRIO EXCLUSIVO DE INVESTIMENTOS
    dados_investimentos = {
        # --- VALORES ATUAIS (O que está na tela) ---
        'fb_atual': st.session_state.get('fb_atual', 0),
        'ig_atual': st.session_state.get('ig_atual', 0),
        'tt_atual': st.session_state.get('tt_atual', 0),
        'google_atual': st.session_state.get('google_atual', 0),
        'yt_atual': st.session_state.get('yt_atual', 0),
        'pmax_atual': st.session_state.get('pmax_atual', 0),
        'total_atual': investimento_total_atual, # Variável que você já calcula no topo do 'if submitted'

        # --- VARIAÇÕES MÊS SOBRE MÊS (MoM - Atual vs Mês Passado) ---
        'var_fb_mes': calcular_variacao(st.session_state.get('fb_atual', 0), st.session_state.get('fb_mes', 0)),
        'var_ig_mes': calcular_variacao(st.session_state.get('ig_atual', 0), st.session_state.get('ig_mes', 0)),
        'var_tt_mes': calcular_variacao(st.session_state.get('tt_atual', 0), st.session_state.get('tt_mes', 0)),
        'var_google_mes': calcular_variacao(st.session_state.get('google_atual', 0), st.session_state.get('google_mes', 0)),
        'var_yt_mes': calcular_variacao(st.session_state.get('yt_atual', 0), st.session_state.get('yt_mes', 0)),
        'var_pmax_mes': calcular_variacao(st.session_state.get('pmax_atual', 0), st.session_state.get('pmax_mes', 0)),
        'var_total_mes': calcular_variacao(investimento_total_atual, investimento_total_mes_passado),

        # --- VARIAÇÕES ANO SOBRE ANO (YoY - Atual vs Ano Passado) ---
        'var_fb_ano': calcular_variacao(st.session_state.get('fb_atual', 0), st.session_state.get('fb_ano', 0)),
        'var_ig_ano': calcular_variacao(st.session_state.get('ig_atual', 0), st.session_state.get('ig_ano', 0)),
        'var_tt_ano': calcular_variacao(st.session_state.get('tt_atual', 0), st.session_state.get('tt_ano', 0)),
        'var_google_ano': calcular_variacao(st.session_state.get('google_atual', 0), st.session_state.get('google_ano', 0)),
        'var_yt_ano': calcular_variacao(st.session_state.get('yt_atual', 0), st.session_state.get('yt_ano', 0)),
        'var_pmax_ano': calcular_variacao(st.session_state.get('pmax_atual', 0), st.session_state.get('pmax_ano', 0)),
        'var_total_ano': calcular_variacao(investimento_total_atual, investimento_total_ano_passado)
    }

    # DICIONÁRIO EXCLUSIVO PARA SEO E CONTEÚDO (COMPLETO)
    dados_seo = {
        # --- TOTAIS (PAGO + ORGÂNICO) ---
        'vis_total_atual': st.session_state.get('seo_vis_atual', 0),
        'vis_total_mes':   st.session_state.get('seo_vis_mes', 0),
        'vis_total_ano':   st.session_state.get('seo_vis_ano', 0),
        
        'sess_total_atual': st.session_state.get('seo_sess_atual', 0),
        'sess_total_mes':   st.session_state.get('seo_sess_mes', 0),
        'sess_total_ano':   st.session_state.get('seo_sess_ano', 0),
        
        'user_total_atual': st.session_state.get('seo_user_atual', 0),
        'user_total_mes':   st.session_state.get('seo_user_mes', 0),
        'user_total_ano':   st.session_state.get('seo_user_ano', 0),

        # --- APENAS ORGÂNICO (SEO PURO) ---
        'vis_org_atual': st.session_state.get('seo_vis_org_atual', 0),
        'vis_org_mes':   st.session_state.get('seo_vis_org_mes', 0),
        'vis_org_ano':   st.session_state.get('seo_vis_org_ano', 0),
        
        'sess_org_atual': st.session_state.get('seo_sess_org_atual', 0),
        'sess_org_mes':   st.session_state.get('seo_sess_org_mes', 0),
        'sess_org_ano':   st.session_state.get('seo_sess_org_ano', 0),
        
        'user_org_atual': st.session_state.get('seo_user_org_atual', 0),
        'user_org_mes':   st.session_state.get('seo_user_org_mes', 0),
        'user_org_ano':   st.session_state.get('seo_user_org_ano', 0),

        # --- VARIAÇÕES MoM (Mês Atual vs Mês Passado) ---
        'var_vis_total_mes': calcular_variacao(st.session_state.get('seo_vis_atual', 0), st.session_state.get('seo_vis_mes', 0)),
        'var_vis_org_mes':   calcular_variacao(st.session_state.get('seo_vis_org_atual', 0), st.session_state.get('seo_vis_org_mes', 0)),
        'var_sess_org_mes':  calcular_variacao(st.session_state.get('seo_sess_org_atual', 0), st.session_state.get('seo_sess_org_mes', 0)),

        # --- VARIAÇÕES YoY (Este Ano vs Ano Passado) ---
        'var_vis_total_ano': calcular_variacao(st.session_state.get('seo_vis_atual', 0), st.session_state.get('seo_vis_ano', 0)),
        'var_vis_org_ano':   calcular_variacao(st.session_state.get('seo_vis_org_atual', 0), st.session_state.get('seo_vis_org_ano', 0)),
        'var_sess_org_ano':  calcular_variacao(st.session_state.get('seo_sess_org_atual', 0), st.session_state.get('seo_sess_org_ano', 0)),

        # --- EXTRAS ---
        'top_keywords': top_keywords,
        'info_concorrentes': info_concorrentes
    }

    # Gerar relatório — 7 etapas sequenciais
    try:
        progress = st.progress(0, text="Iniciando pipeline de inteligência (7 etapas)...")

        with st.spinner("1/7 — Cenário Atual..."):
            etapa_cenario_atual = gerar_cenario_atual(dados_metrica_performance, dados_investimentos, dados_custos, info_concorrentes, modelo_escolhido)
        progress.progress(1/7, text="1/7 — Cenário Atual")

        with st.spinner("2/7 — Destaques..."):
            etapa_destaques = gerar_destaques(etapa_cenario_atual, modelo_escolhido)
        progress.progress(2/7, text="2/7 — Destaques")

        with st.spinner("3/7 — Mídias Pagas..."):
            etapa_midias_pagas = gerar_midias_pagas(etapa_cenario_atual, dados_investimentos, dados_custos, modelo_escolhido)
        progress.progress(3/7, text="3/7 — Mídias Pagas")

        with st.spinner("4/7 — Social..."):
            etapa_social = gerar_social(etapa_cenario_atual, descricoes_imagens, descricoes_imagens_mes_passado, dados_custos, resumos_social_csvs, modelo_escolhido)
        progress.progress(4/7, text="4/7 — Social")

        with st.spinner("5/7 — SEO..."):
            etapa_seo = gerar_seo_content(etapa_cenario_atual, dados_seo, dados_custos, resumos_seo_csvs, modelo_escolhido)
        progress.progress(5/7, text="5/7 — SEO")

        with st.spinner("6/7 — Aprendizados..."):
            etapa_aprendizados = gerar_aprendizados(etapa_cenario_atual, etapa_destaques, etapa_midias_pagas, etapa_social, etapa_seo, dados_metrica_performance, dados_custos, dados_seo, modelo_escolhido)
        progress.progress(6/7, text="6/7 — Aprendizados")

        with st.spinner("7/7 — Próximos Passos..."):
            etapa_proximos_passos = gerar_proximos_passos(etapa_cenario_atual, etapa_aprendizados, modelo_escolhido)
        progress.progress(7/7, text="7/7 — Pipeline completo!")

        # Armazenar resultados
        st.session_state.relatorio_gerado = True
        st.session_state.dados_processados = dados_metrica_performance
        st.session_state.descricoes_imagens = descricoes_imagens
        st.session_state.descricoes_imagens_mes_passado = descricoes_imagens_mes_passado
        st.session_state.resumos_social_csvs = resumos_social_csvs
        st.session_state.resumos_seo_csvs = resumos_seo_csvs
        st.session_state.etapa_cenario_atual = etapa_cenario_atual
        st.session_state.etapa_destaques = etapa_destaques
        st.session_state.etapa_midias_pagas = etapa_midias_pagas
        st.session_state.etapa_social = etapa_social
        st.session_state.etapa_seo = etapa_seo
        st.session_state.etapa_aprendizados = etapa_aprendizados
        st.session_state.etapa_proximos_passos = etapa_proximos_passos

        st.rerun()

    except Exception as e:
        st.error(f"Erro ao gerar relatório: {str(e)}")

# Exibir relatório
if st.session_state.relatorio_gerado:
    st.markdown("---")
    st.header("📄 Relatório Executivo Gerado")

    dados = st.session_state.dados_processados

    # =====================================================================
    # PAINEL DE CARDS — KPIs PRINCIPAIS
    # =====================================================================
    st.subheader("📊 Painel de Performance")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Investimento", f"R$ {dados.get('spend_atual', 0):,.0f}",
                  f"{dados.get('var_invest_mes', 0):+.1f}% MoM")
    with col2:
        st.metric("Cliques", f"{dados.get('cli_atual', 0):,}",
                  f"{dados.get('var_cli_mes', 0):+.1f}% MoM")
    with col3:
        st.metric("Impressões", f"{dados.get('imp_atual', 0):,}",
                  f"{dados.get('var_imp_mes', 0):+.1f}% MoM")
    with col4:
        st.metric("CTR", f"{dados.get('ctr_atual', 0):.2f}%",
                  f"{dados.get('var_ctr_mes', 0):+.1f}% MoM")

    col1, col2, col3, col4 = st.columns(4)
    with col1:
        st.metric("Alcance", f"{dados.get('reach_atual', 0):,}",
                  f"{dados.get('var_reach_mes', 0):+.1f}% MoM")
    with col2:
        st.metric("Sessões", f"{dados.get('sess_atual', 0):,}",
                  f"{dados.get('var_sess_mes', 0):+.1f}% MoM")
    with col3:
        st.metric("Engajamentos", f"{dados.get('eng_atual', 0):,}",
                  f"{dados.get('var_eng_mes', 0):+.1f}% MoM")
    with col4:
        st.metric("Video Thruplays", f"{dados.get('vtp_atual', 0):,}",
                  f"{dados.get('var_vtp_mes', 0):+.1f}% MoM")

    st.markdown("---")

    # =====================================================================
    # GRÁFICOS — Linha 1: Variações MoM vs YoY (Barras agrupadas)
    # =====================================================================
    st.subheader("📈 Variações MoM vs YoY")

    metricas_nomes = ["Invest.", "Sessões", "Alcance", "Impressões", "Cliques", "Engajam.", "CTR"]
    var_mom = [
        dados.get('var_invest_mes', 0), dados.get('var_sess_mes', 0),
        dados.get('var_reach_mes', 0), dados.get('var_imp_mes', 0),
        dados.get('var_cli_mes', 0), dados.get('var_eng_mes', 0),
        dados.get('var_ctr_mes', 0)
    ]
    var_yoy = [
        dados.get('var_invest_ano', 0), dados.get('var_sess_ano', 0),
        dados.get('var_reach_ano', 0), dados.get('var_imp_ano', 0),
        dados.get('var_cli_ano', 0), dados.get('var_eng_ano', 0),
        dados.get('var_ctr_ano', 0)
    ]

    fig_variacoes = go.Figure()
    fig_variacoes.add_trace(go.Bar(
        name='MoM (%)', x=metricas_nomes, y=var_mom,
        marker_color=['#27AE60' if v >= 0 else '#E74C3C' for v in var_mom],
        text=[f"{v:+.1f}%" for v in var_mom], textposition='outside',
        textfont=dict(size=10)
    ))
    fig_variacoes.add_trace(go.Bar(
        name='YoY (%)', x=metricas_nomes, y=var_yoy,
        marker_color=['#2E86C1' if v >= 0 else '#E67E22' for v in var_yoy],
        text=[f"{v:+.1f}%" for v in var_yoy], textposition='outside',
        textfont=dict(size=10)
    ))
    fig_variacoes.update_layout(
        barmode='group', height=400,
        plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)',
        font=dict(family='Calibri', size=12),
        legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
        yaxis=dict(title='Variação (%)', zeroline=True, zerolinecolor='#BDC3C7', gridcolor='#EAECEE'),
        xaxis=dict(title=''),
        margin=dict(t=40, b=40)
    )
    st.plotly_chart(fig_variacoes, use_container_width=True)

    # =====================================================================
    # GRÁFICOS — Linha 2: Investimento por Canal (Pizza) + Custos (Radar)
    # =====================================================================
    col_g1, col_g2 = st.columns(2)

    with col_g1:
        st.markdown("**Distribuição de Investimento por Canal**")
        canais = ['Facebook', 'Instagram', 'TikTok', 'Google Ads', 'YouTube', 'PMax']
        valores_canais = [
            st.session_state.get('fb_atual', 0), st.session_state.get('ig_atual', 0),
            st.session_state.get('tt_atual', 0), st.session_state.get('google_atual', 0),
            st.session_state.get('yt_atual', 0), st.session_state.get('pmax_atual', 0)
        ]
        # Filtrar canais com valor > 0
        canais_filtrados = [c for c, v in zip(canais, valores_canais) if v > 0]
        valores_filtrados = [v for v in valores_canais if v > 0]

        if valores_filtrados:
            fig_pizza = go.Figure(data=[go.Pie(
                labels=canais_filtrados, values=valores_filtrados,
                hole=0.45,
                marker=dict(colors=['#3498DB', '#E91E63', '#000000', '#4CAF50', '#FF0000', '#FF9800']),
                textinfo='label+percent', textposition='outside',
                textfont=dict(size=11)
            )])
            fig_pizza.update_layout(
                height=380, showlegend=False,
                margin=dict(t=20, b=20, l=20, r=20),
                annotations=[dict(text=f"R$ {sum(valores_filtrados):,.0f}", x=0.5, y=0.5,
                                  font_size=14, font_color='#1B3A5C', showarrow=False)]
            )
            st.plotly_chart(fig_pizza, use_container_width=True)
        else:
            st.info("Nenhum investimento por canal informado.")

    with col_g2:
        st.markdown("**Indicadores de Custo (Atual vs Mês Passado)**")
        custos_labels = ['CPC', 'CPM', 'CPE', 'CPV']
        custos_atual = [
            st.session_state.get('cpc_atual', 0), st.session_state.get('cpm_atual', 0),
            st.session_state.get('cpe_atual', 0), st.session_state.get('cpv_atual', 0)
        ]
        custos_mes = [
            st.session_state.get('cpc_mes', 0), st.session_state.get('cpm_mes', 0),
            st.session_state.get('cpe_mes', 0), st.session_state.get('cpv_mes', 0)
        ]

        if any(v > 0 for v in custos_atual):
            fig_custos = go.Figure()
            fig_custos.add_trace(go.Bar(
                name='Atual', x=custos_labels, y=custos_atual,
                marker_color='#1B3A5C',
                text=[f"R$ {v:.2f}" for v in custos_atual], textposition='outside'
            ))
            fig_custos.add_trace(go.Bar(
                name='Mês Passado', x=custos_labels, y=custos_mes,
                marker_color='#AEB6BF',
                text=[f"R$ {v:.2f}" for v in custos_mes], textposition='outside'
            ))
            fig_custos.update_layout(
                barmode='group', height=380,
                plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)',
                font=dict(family='Calibri', size=12),
                legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
                yaxis=dict(title='R$', gridcolor='#EAECEE'),
                margin=dict(t=40, b=20)
            )
            st.plotly_chart(fig_custos, use_container_width=True)
        else:
            st.info("Nenhum indicador de custo informado.")

    # =====================================================================
    # GRÁFICO — Efeito Tesoura (se aplicável)
    # =====================================================================
    var_inv_mom = dados.get('var_invest_mes', 0)
    var_cli_mom = dados.get('var_cli_mes', 0)
    var_eng_mom = dados.get('var_eng_mes', 0)
    tesoura_detectada = (var_inv_mom < -5 and (var_cli_mom > 0 or var_eng_mom > 0))

    if tesoura_detectada:
        st.subheader("✂️ Efeito Tesoura Detectado")
        st.caption("Investimento em queda com resultados em alta — sinal de ganho de produtividade digital.")

        periodos = ['Mês Passado', 'Mês Atual']
        inv_vals = [100, 100 + var_inv_mom]
        cli_vals = [100, 100 + var_cli_mom]

        fig_tesoura = go.Figure()
        fig_tesoura.add_trace(go.Scatter(
            x=periodos, y=inv_vals, name='Investimento',
            line=dict(color='#E74C3C', width=3, dash='dash'),
            mode='lines+markers+text',
            text=[f"{inv_vals[0]:.0f}", f"{inv_vals[1]:.0f}"], textposition='top center'
        ))
        fig_tesoura.add_trace(go.Scatter(
            x=periodos, y=cli_vals, name='Cliques',
            line=dict(color='#27AE60', width=3),
            mode='lines+markers+text',
            text=[f"{cli_vals[0]:.0f}", f"{cli_vals[1]:.0f}"], textposition='bottom center'
        ))
        fig_tesoura.update_layout(
            height=300,
            plot_bgcolor='rgba(0,0,0,0)', paper_bgcolor='rgba(0,0,0,0)',
            font=dict(family='Calibri'), yaxis=dict(title='Índice (base 100)', gridcolor='#EAECEE'),
            legend=dict(orientation='h', yanchor='bottom', y=1.02, xanchor='right', x=1),
            margin=dict(t=40, b=30)
        )
        st.plotly_chart(fig_tesoura, use_container_width=True)

    # =====================================================================
    # TABELA COMPARATIVA COMPLETA
    # =====================================================================
    st.subheader("📋 Tabela Comparativa Completa")
    tabela_dados = {
        'Métrica': ['Investimento', 'Sessões', 'Alcance', 'Impressões', 'Cliques', 'Engajamentos', 'CTR (%)'],
        'Atual': [
            f"R$ {dados.get('spend_atual', 0):,.0f}", f"{dados.get('sess_atual', 0):,}",
            f"{dados.get('reach_atual', 0):,}", f"{dados.get('imp_atual', 0):,}",
            f"{dados.get('cli_atual', 0):,}", f"{dados.get('eng_atual', 0):,}",
            f"{dados.get('ctr_atual', 0):.2f}%"
        ],
        'Var. MoM': [
            f"{dados.get('var_invest_mes', 0):+.1f}%", f"{dados.get('var_sess_mes', 0):+.1f}%",
            f"{dados.get('var_reach_mes', 0):+.1f}%", f"{dados.get('var_imp_mes', 0):+.1f}%",
            f"{dados.get('var_cli_mes', 0):+.1f}%", f"{dados.get('var_eng_mes', 0):+.1f}%",
            f"{dados.get('var_ctr_mes', 0):+.1f}%"
        ],
        'Var. YoY': [
            f"{dados.get('var_invest_ano', 0):+.1f}%", f"{dados.get('var_sess_ano', 0):+.1f}%",
            f"{dados.get('var_reach_ano', 0):+.1f}%", f"{dados.get('var_imp_ano', 0):+.1f}%",
            f"{dados.get('var_cli_ano', 0):+.1f}%", f"{dados.get('var_eng_ano', 0):+.1f}%",
            f"{dados.get('var_ctr_ano', 0):+.1f}%"
        ]
    }
    st.dataframe(pd.DataFrame(tabela_dados), use_container_width=True, hide_index=True)

    st.markdown("---")

    # =====================================================================
    # SEÇÕES ANALÍTICAS DO RELATÓRIO (Pipeline de 7 etapas)
    # =====================================================================
    st.subheader("📌 1 — Cenário Atual")
    st.write(st.session_state.etapa_cenario_atual)

    st.subheader("⭐ 2 — Destaques")
    st.write(st.session_state.etapa_destaques)

    st.subheader("💰 3 — Mídias Pagas")
    st.write(st.session_state.etapa_midias_pagas)

    st.subheader("📱 4 — Social")
    st.write(st.session_state.etapa_social)

    st.subheader("🔍 5 — SEO")
    st.write(st.session_state.etapa_seo)

    st.subheader("💡 6 — Aprendizados")
    st.write(st.session_state.etapa_aprendizados)

    st.subheader("🚀 7 — Próximos Passos")
    st.write(st.session_state.etapa_proximos_passos)
    
    # =====================================================================
    # DOWNLOADS — 4 DOCUMENTOS
    # =====================================================================
    st.markdown("---")
    st.subheader("📥 Documentos para Download")

    try:
        # Montar dicionários de apoio para DOCX
        dados_inv_docx = {
            'fb_atual': st.session_state.get('fb_atual', 0),
            'ig_atual': st.session_state.get('ig_atual', 0),
            'tt_atual': st.session_state.get('tt_atual', 0),
            'google_atual': st.session_state.get('google_atual', 0),
            'yt_atual': st.session_state.get('yt_atual', 0),
            'pmax_atual': st.session_state.get('pmax_atual', 0),
            'total_atual': (st.session_state.get('fb_atual', 0) + st.session_state.get('ig_atual', 0) +
                            st.session_state.get('tt_atual', 0) + st.session_state.get('google_atual', 0) +
                            st.session_state.get('yt_atual', 0) + st.session_state.get('pmax_atual', 0)),
            'var_fb_mes': calcular_variacao(st.session_state.get('fb_atual', 0), st.session_state.get('fb_mes', 0)),
            'var_ig_mes': calcular_variacao(st.session_state.get('ig_atual', 0), st.session_state.get('ig_mes', 0)),
            'var_tt_mes': calcular_variacao(st.session_state.get('tt_atual', 0), st.session_state.get('tt_mes', 0)),
            'var_google_mes': calcular_variacao(st.session_state.get('google_atual', 0), st.session_state.get('google_mes', 0)),
            'var_total_mes': 0,
            'var_fb_ano': calcular_variacao(st.session_state.get('fb_atual', 0), st.session_state.get('fb_ano', 0)),
            'var_ig_ano': calcular_variacao(st.session_state.get('ig_atual', 0), st.session_state.get('ig_ano', 0)),
            'var_tt_ano': calcular_variacao(st.session_state.get('tt_atual', 0), st.session_state.get('tt_ano', 0)),
            'var_google_ano': calcular_variacao(st.session_state.get('google_atual', 0), st.session_state.get('google_ano', 0)),
            'var_total_ano': 0,
        }
        dados_custos_docx = {
            'cpc_atual': st.session_state.get('cpc_atual', 0), 'cpm_atual': st.session_state.get('cpm_atual', 0),
            'cpe_atual': st.session_state.get('cpe_atual', 0), 'cpv_atual': st.session_state.get('cpv_atual', 0),
            'var_cpc_mes': calcular_variacao(st.session_state.get('cpc_atual', 0), st.session_state.get('cpc_mes', 0)),
            'var_cpm_mes': calcular_variacao(st.session_state.get('cpm_atual', 0), st.session_state.get('cpm_mes', 0)),
            'var_cpe_mes': calcular_variacao(st.session_state.get('cpe_atual', 0), st.session_state.get('cpe_mes', 0)),
            'var_cpv_mes': calcular_variacao(st.session_state.get('cpv_atual', 0), st.session_state.get('cpv_mes', 0)),
            'var_cpc_ano': calcular_variacao(st.session_state.get('cpc_atual', 0), st.session_state.get('cpc_ano', 0)),
            'var_cpm_ano': calcular_variacao(st.session_state.get('cpm_atual', 0), st.session_state.get('cpm_ano', 0)),
            'var_cpe_ano': calcular_variacao(st.session_state.get('cpe_atual', 0), st.session_state.get('cpe_ano', 0)),
            'var_cpv_ano': calcular_variacao(st.session_state.get('cpv_atual', 0), st.session_state.get('cpv_ano', 0)),
        }
        dados_seo_docx = {
            'vis_total_atual': st.session_state.get('seo_vis_atual', 0),
            'sess_org_atual': st.session_state.get('seo_sess_org_atual', 0),
            'vis_org_atual': st.session_state.get('seo_vis_org_atual', 0),
            'var_vis_total_mes': calcular_variacao(st.session_state.get('seo_vis_atual', 0), st.session_state.get('seo_vis_mes', 0)),
            'var_sess_org_mes': calcular_variacao(st.session_state.get('seo_sess_org_atual', 0), st.session_state.get('seo_sess_org_mes', 0)),
            'var_vis_org_mes': calcular_variacao(st.session_state.get('seo_vis_org_atual', 0), st.session_state.get('seo_vis_org_mes', 0)),
            'var_vis_total_ano': calcular_variacao(st.session_state.get('seo_vis_atual', 0), st.session_state.get('seo_vis_ano', 0)),
            'var_sess_org_ano': calcular_variacao(st.session_state.get('seo_sess_org_atual', 0), st.session_state.get('seo_sess_org_ano', 0)),
            'var_vis_org_ano': calcular_variacao(st.session_state.get('seo_vis_org_atual', 0), st.session_state.get('seo_vis_org_ano', 0)),
        }

        ts = datetime.now().strftime('%Y%m%d_%H%M')

        # Coletar etapas do session state
        etapas = {
            'etapa_cenario_atual': st.session_state.get('etapa_cenario_atual', ''),
            'etapa_destaques': st.session_state.get('etapa_destaques', ''),
            'etapa_midias_pagas': st.session_state.get('etapa_midias_pagas', ''),
            'etapa_social': st.session_state.get('etapa_social', ''),
            'etapa_seo': st.session_state.get('etapa_seo', ''),
            'etapa_aprendizados': st.session_state.get('etapa_aprendizados', ''),
            'etapa_proximos_passos': st.session_state.get('etapa_proximos_passos', ''),
        }

        # --- DOC 1: RELATÓRIO INTERNO (DOCX) ---
        docx_interno = gerar_docx_relatorio(
            dados=dados, dados_investimentos=dados_inv_docx, dados_custos=dados_custos_docx,
            dados_seo=dados_seo_docx, **etapas,
        )

        # --- DOC 2: RELATÓRIO PARA O CLIENTE (DOCX) ---
        docx_cliente = gerar_docx_cliente(
            dados=dados, dados_investimentos=dados_inv_docx,
            dados_custos=dados_custos_docx, dados_seo=dados_seo_docx, **etapas,
        )

        # Layout de 2 botões
        col_d1, col_d2 = st.columns(2)
        with col_d1:
            st.download_button(
                label="📄 Relatório Interno (DOCX)",
                data=docx_interno,
                file_name=f"INTERNO_relatorio_executivo_{ts}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Relatório completo para uso interno da Macfor — 7 seções com análise profunda."
            )
        with col_d2:
            st.download_button(
                label="📄 Relatório para o Cliente (DOCX)",
                data=docx_cliente,
                file_name=f"CLIENTE_relatorio_resultados_syngenta_{ts}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                help="Relatório de resultados para apresentação ao cliente Syngenta."
            )

    except Exception as e:
        st.error(f"Erro ao gerar documentos: {str(e)}")
